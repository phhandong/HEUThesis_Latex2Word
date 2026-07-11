from __future__ import annotations

import re
import base64
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .bibliography import ReferenceEntry
from .latex_parser import ThesisMetadata
from .preprocess import (
    CITATION_PLACEHOLDER_RE,
    EQUATION_MARKER_RE,
    EQUATION_REF_PLACEHOLDER_RE,
    HEU_COVER_PLACEHOLDER,
)
from .report import ConversionReport


A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
BODY_SIZE_PT = 12.0
BODY_LINE_SPACING_PT = 20.5
FIRST_LINE_INDENT_PT = 24.0
CHAPTER_SIZE_PT = 18.0
CHAPTER_LINE_SPACING_PT = 28.5
CHAPTER_SPACE_BEFORE_PT = 28.35
CHAPTER_SPACE_AFTER_PT = 28.75
SECTION_SIZE_PT = 15.0
SECTION_LINE_SPACING_PT = 21.0
SECTION_SPACE_PT = 19.84
SUBSECTION_SIZE_PT = 14.0
SUBSECTION_LINE_SPACING_PT = 18.0
SUBSECTION_SPACE_PT = 17.01
SUBSUBSECTION_SIZE_PT = 12.0
SUBSUBSECTION_LINE_SPACING_PT = 20.5
SUBSUBSECTION_SPACE_PT = 8.5
CAPTION_SIZE_PT = 10.5
CAPTION_LINE_SPACING_PT = 13.65
TABLE_SIZE_PT = 10.5
TABLE_LINE_SPACING_PT = 13.65

FRONT_MATTER_TITLES = {
    "摘要",
    "Abstract",
    "目录",
    "参考文献",
    "学位论文原创性声明",
    "学位论文授权使用声明",
    "学位论文原创性声明和授权使用声明",
    "插图清单",
    "附表清单",
}
ROMAN_PAGE_TITLES = {
    "摘要",
    "Abstract",
    "目录",
    "插图清单",
    "附表清单",
}

CAPTION_PREFIX_RE = re.compile(r"^[图表]\s*\d+(?:\.\d+)+")
REFERENCE_ENTRY_RE = re.compile(r"^\[\d+\]")
CHAPTER_PREFIX_RE = re.compile(r"^第\s*([0-9一二三四五六七八九十百]+)\s*章\s*")
COVER_MARKER = "HEU_COVER_PLACEHOLDER"
HEU_COVER_PLACEHOLDER = COVER_MARKER
DECLARATION_MARKER = "HEU_DECLARATION_PLACEHOLDER"
PUBLICATION_SECTION_TITLES = {
    "（一）发表的相关论文",
    "（二）申请及已获得的专利（无专利时此项不必列出）",
    "（三）参与的科研项目及获奖情况",
}
PRIMARY_BACK_MATTER_TITLES = {
    "结论",
    "致谢",
    "攻读硕士学位期间发表的论文和取得的科研成果",
}


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _engineering_template_path(filename: str) -> Path:
    return _repo_root() / "requirements" / "Engineering" / filename


COVER_TEMPLATE_PATH = _repo_root() / "cover.docx"
COVER_FALLBACK_TEMPLATE_PATH = _engineering_template_path("封面.docx")
DECLARATION_TEMPLATE_PATH = _engineering_template_path("5.学位论文原创性声明和授权使用声明(最新).docx")
HEU_STYLES = {
    "body": "正文",
    "body_inline_math": "正文（内嵌公式）",
    "chapter": "章标题",
    "section": "节标题",
    "subsection": "小节标题",
    "subsubsection": "四级标题",
    "caption": "图表题注",
    "figure": "图片",
    "display_math": "独立公式",
    "reference": "参考文献",
    "cover_title": "封面题名",
    "cover_meta": "封面信息",
    "cover_small": "封面小字",
}
TOC_STYLE_NAMES = ("toc 1", "toc 2", "toc 3")
LEGACY_HEU_STYLE_NAMES = {
    "HEU Body",
    "HEU Body Inline Math",
    "HEU Chapter Title",
    "HEU Section Title",
    "HEU Subsection Title",
    "HEU Subsubsection Title",
    "HEU Caption",
    "HEU Figure",
    "HEU Display Math",
    "HEU Reference",
    "HEU Cover Title",
    "HEU Cover Meta",
    "HEU Cover Small",
}
KEEP_CUSTOM_STYLE_NAMES = set(HEU_STYLES.values()) | {
    "封面宋体小2",
    "封面2号黑体英文论文名称",
    "英文扉页姓名信息",
    "原创性声明内容样式",
} | set(TOC_STYLE_NAMES)


def _set_run_font(run, font_cn: str, font_en: str, size_pt: float, bold: bool | None = None) -> None:
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def _set_style_font(style, font_cn: str, font_en: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = font_en
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def _remove_child(element, tag: str) -> None:
    if element is None:
        return
    for child in list(element.findall(qn(tag))):
        element.remove(child)


def _clear_text_direct_formatting(paragraph) -> None:
    ppr = paragraph._p.pPr
    if ppr is not None:
        for tag in ("w:jc", "w:spacing", "w:ind"):
            _remove_child(ppr, tag)
    for run in paragraph.runs:
        rpr = run._r.rPr
        if rpr is None:
            continue
        for tag in ("w:rStyle", "w:rFonts", "w:sz", "w:szCs", "w:color"):
            _remove_child(rpr, tag)


def _set_paragraph_layout(
    paragraph,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None,
    first_line_indent_pt: float | None,
    line_spacing_pt: float,
    space_before_pt: float,
    space_after_pt: float,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = (
        None if first_line_indent_pt is None else Pt(first_line_indent_pt)
    )
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)


def _set_style_paragraph_layout(
    style,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None,
    first_line_indent_pt: float | None,
    line_spacing_pt: float,
    space_before_pt: float,
    space_after_pt: float,
) -> None:
    if alignment is not None:
        style.paragraph_format.alignment = alignment
    style.paragraph_format.first_line_indent = (
        None if first_line_indent_pt is None else Pt(first_line_indent_pt)
    )
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(line_spacing_pt)
    style.paragraph_format.space_before = Pt(space_before_pt)
    style.paragraph_format.space_after = Pt(space_after_pt)


def _set_style_single_line_layout(
    style,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None,
    first_line_indent_pt: float | None,
    space_before_pt: float,
    space_after_pt: float,
) -> None:
    if alignment is not None:
        style.paragraph_format.alignment = alignment
    style.paragraph_format.first_line_indent = (
        None if first_line_indent_pt is None else Pt(first_line_indent_pt)
    )
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(space_before_pt)
    style.paragraph_format.space_after = Pt(space_after_pt)


def _set_style_exact_line_layout(
    style,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None,
    first_line_indent_pt: float | None,
    line_spacing_pt: float,
    space_before_pt: float,
    space_after_pt: float,
) -> None:
    if alignment is not None:
        style.paragraph_format.alignment = alignment
    style.paragraph_format.first_line_indent = (
        None if first_line_indent_pt is None else Pt(first_line_indent_pt)
    )
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(line_spacing_pt)
    style.paragraph_format.space_before = Pt(space_before_pt)
    style.paragraph_format.space_after = Pt(space_after_pt)


def _add_field(paragraph, instruction: str, fallback: str = "1"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = fallback
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    return run


def _add_ref_field(paragraph, bookmark: str, fallback: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' REF {bookmark} \\n \\# "0" \\h '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = fallback
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    return run


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _get_or_add_paragraph_properties(element):
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        element.insert(0, ppr)
    return ppr


def _set_keep_next(element, enabled: bool = True) -> None:
    ppr = _get_or_add_paragraph_properties(element)
    _remove_child(ppr, "w:keepNext")
    if enabled:
        ppr.append(OxmlElement("w:keepNext"))


def _set_keep_lines(element, enabled: bool = True) -> None:
    ppr = _get_or_add_paragraph_properties(element)
    _remove_child(ppr, "w:keepLines")
    if enabled:
        ppr.append(OxmlElement("w:keepLines"))


def _reset_cover_field_indent(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for old in list(ppr.findall(qn("w:ind"))):
        ppr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:firstLine"), "0")
    ppr.append(ind)


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_cell_text(cell, text: str) -> None:
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _set_paragraph_text(para, text)
    for extra in cell.paragraphs[1:]:
        _clear_paragraph(extra)


def _strip_cover_annotations(text: str) -> str:
    text = re.sub(r"（[^（）]*(?:宋体|黑体|楷体|Times New Roman)[^（）]*）", "", text)
    return re.sub(r"[ \t]+$", "", text)


def _set_cover_top_line(
    paragraph,
    left_label: str,
    left_value: str,
    right_label: str,
    right_value: str,
) -> None:
    _clear_paragraph(paragraph)
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run(left_label)
    left_run = paragraph.add_run(left_value or "               ")
    left_run.underline = True
    paragraph.add_run("\t")
    paragraph.add_run(right_label)
    right_run = paragraph.add_run(right_value or "               ")
    right_run.underline = True


def _set_cover_field_line(paragraph, label: str, value: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.tab_stops.clear_all()
    _clear_paragraph(paragraph)
    _reset_cover_field_indent(paragraph)
    run = paragraph.add_run(f"　　　　　　　　　{_spread_cover_label(label)}{value}")
    _set_run_font(run, "宋体", "Times New Roman", 15.0, False)


def _set_cover_plain_text(paragraph, text: str) -> None:
    _set_paragraph_text(paragraph, _strip_cover_annotations(text))


def _spread_cover_label(label: str) -> str:
    label = label.replace("：", "").replace(" ", "")
    label_map = {
        "硕士研究生": "硕士研究生：　",
        "博士研究生": "博士研究生：　",
        "指导教师": "指导教师：　　",
        "校外导师": "校外导师：　　",
        "学位类别": "学位类别：　　",
    }
    if label in label_map:
        return label_map[label]
    if label == "学位类别":
        return "学位类别：　　"
    if len(label) <= 1:
        return f"{label}："
    return f"{' '.join(label)}："


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _section_properties_from_doc(doc: Document):
    for child in reversed(list(doc._element.body)):
        if child.tag == qn("w:sectPr"):
            return child
    return None


def _set_section_start_type(sect_pr, section_type: str) -> None:
    for old in list(sect_pr.findall(qn("w:type"))):
        sect_pr.remove(old)
    if section_type:
        sect_type = OxmlElement("w:type")
        sect_type.set(qn("w:val"), section_type)
        sect_pr.insert(0, sect_type)


def _strip_header_footer_references(element) -> None:
    nodes = [element] if element.tag == qn("w:sectPr") else list(element.iter(qn("w:sectPr")))
    for sect_pr in nodes:
        for tag in ("w:headerReference", "w:footerReference"):
            for ref in list(sect_pr.findall(qn(tag))):
                sect_pr.remove(ref)


def _section_break_paragraph(doc: Document, section_type: str = "nextPage"):
    source = _section_properties_from_doc(doc)
    sect_pr = deepcopy(source) if source is not None else OxmlElement("w:sectPr")
    _strip_header_footer_references(sect_pr)
    _set_section_start_type(sect_pr, section_type)
    paragraph = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    paragraph.append(ppr)
    ppr.append(sect_pr)
    return paragraph


def _paragraph_section_properties(element):
    if element.tag != qn("w:p"):
        return None
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        return None
    return ppr.find(qn("w:sectPr"))


def _set_section_page_numbering(sect_pr, fmt: str | None, start: int | None = None) -> None:
    for old in list(sect_pr.findall(qn("w:pgNumType"))):
        sect_pr.remove(old)
    if fmt is None:
        return
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    sect_pr.append(pg_num)


def _block_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def _section_blocks(doc: Document):
    blocks = []
    for child in doc._element.body:
        if child.tag == qn("w:sectPr"):
            if blocks:
                yield blocks, child
                blocks = []
            continue
        blocks.append(child)
        sect_pr = _paragraph_section_properties(child)
        if sect_pr is not None:
            yield blocks, sect_pr
            blocks = []


def _first_meaningful_section_text(blocks) -> str:
    for block in blocks:
        text = _block_text(block)
        if text:
            return text
    return ""


def _apply_section_page_numbering_by_content(doc: Document) -> None:
    decimal_started = False
    for blocks, sect_pr in _section_blocks(doc):
        start_text = _first_meaningful_section_text(blocks)
        _set_section_page_numbering(sect_pr, None)
        if start_text in ROMAN_PAGE_TITLES:
            _set_section_page_numbering(sect_pr, "upperRoman", 1 if start_text == "摘要" else None)
        elif CHAPTER_PREFIX_RE.match(start_text):
            _set_section_page_numbering(sect_pr, "decimal", 1 if not decimal_started else None)
            decimal_started = True
        elif decimal_started:
            _set_section_page_numbering(sect_pr, "decimal", None)


def _page_break_paragraph():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def _insert_element_before(paragraph, element) -> None:
    parent = paragraph._p.getparent()
    parent.insert(parent.index(paragraph._p), element)


def _insert_section_break_before(paragraph, section_type: str = "nextPage") -> None:
    parent_doc = paragraph._parent.part.document
    _insert_element_before(paragraph, _section_break_paragraph(parent_doc, section_type))


def _ensure_section_break_before(paragraph, section_type: str = "nextPage"):
    sect_pr = _previous_section_properties(paragraph)
    if sect_pr is not None:
        _set_section_start_type(sect_pr, section_type)
        return sect_pr
    parent_doc = paragraph._parent.part.document
    element = _section_break_paragraph(parent_doc, section_type)
    _insert_element_before(paragraph, element)
    return _paragraph_section_properties(element)


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    parent.remove(paragraph._p)


def _is_empty_section_break_paragraph(element) -> bool:
    return element.tag == qn("w:p") and not _element_text(element).strip() and element.find(qn("w:pPr")) is not None and element.find(".//" + qn("w:sectPr")) is not None


def _has_previous_section_break(paragraph) -> bool:
    parent = paragraph._p.getparent()
    index = parent.index(paragraph._p)
    for previous in reversed(parent[:index]):
        if previous.tag != qn("w:p"):
            return False
        if _element_text(previous).strip():
            return False
        if _is_empty_section_break_paragraph(previous):
            return True
    return False


def _previous_section_properties(paragraph):
    parent = paragraph._p.getparent()
    index = parent.index(paragraph._p)
    for previous in reversed(parent[:index]):
        if previous.tag != qn("w:p"):
            return None
        if _element_text(previous).strip():
            return None
        sect_pr = _paragraph_section_properties(previous)
        if sect_pr is not None:
            return sect_pr
    return None


def _iter_all_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _split_cover_title(title: str) -> list[str]:
    lines = [part.strip() for part in re.split(r"[\r\n]+", title) if part.strip()]
    return lines or [title.strip()] if title.strip() else [""]


def _supervisor_cn(metadata: ThesisMetadata) -> str:
    supervisor = metadata.supervisor_cn
    title = metadata.professional_title_cn
    if supervisor and title and title not in supervisor:
        return f"{supervisor} {title}"
    return supervisor


def _cover_table_label(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _english_cover_table_values(metadata: ThesisMetadata) -> dict[str, str]:
    return {
        "Candidate :": metadata.author_en,
        "Supervisor :": metadata.supervisor_en,
        "Associate Supervisor :": metadata.cosupervisor_en,
        "Professional category :": metadata.subject_en,
        "College :": metadata.affiliation_en,
        "Date of Submission :": metadata.submit_date_en,
        "Date of Oral Examination :": metadata.oral_date_en,
        "University :": "Harbin Engineering University",
    }


def _apply_english_cover_paragraphs(doc: Document, metadata: ThesisMetadata) -> None:
    title_lines = _split_cover_title(metadata.title_en)
    title_paragraphs = [
        para
        for para in doc.paragraphs
        if (para.style.name if para.style else "") == "封面2号黑体英文论文名称"
    ]
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("Classified Index"):
            _set_cover_plain_text(para, f"Classified Index: {metadata.classified_index}")
        elif text.startswith("U.D.C"):
            _set_cover_plain_text(para, f"U.D.C: {metadata.udc}")
        elif text.startswith("A Thesis for the Degree of"):
            _set_cover_plain_text(para, f"A Thesis for the Degree of {metadata.degree_label_en}")
    for index, para in enumerate(title_paragraphs):
        if index == len(title_paragraphs) - 1:
            value = " ".join(title_lines[index:]) if index < len(title_lines) else ""
        else:
            value = title_lines[index] if index < len(title_lines) else ""
        _set_cover_plain_text(para, value)


def _apply_cover_metadata(doc: Document, metadata: ThesisMetadata) -> None:
    title_cn = metadata.display_title_cn
    school = "哈尔滨工程大学"
    supervisor_cn = _supervisor_cn(metadata)
    english_table_values = _english_cover_table_values(metadata)

    for para in _iter_all_paragraphs(doc):
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("分类号："):
            _set_cover_top_line(
                para,
                "分类号：",
                metadata.classified_index,
                "密级：",
                metadata.state_secrets,
            )
        elif text.startswith("U D C"):
            _set_cover_top_line(
                para,
                "U D C ：",
                metadata.udc,
                "编号：",
                metadata.document_number,
            )
        elif "专业学位硕士学位论文" in text or "专业学位博士学位论文" in text:
            _set_cover_plain_text(para, metadata.document_label)
        elif "基于FPGA的高能效XXX应用系统" in text:
            _set_cover_plain_text(para, title_cn)
        elif text.startswith("硕士研究生：") or text.startswith("博士研究生："):
            _set_cover_field_line(para, metadata.student_label_cn, metadata.author_cn)
        elif text.startswith("指导教师："):
            _set_cover_field_line(para, "指导教师", supervisor_cn)
        elif text.startswith("校外导师："):
            _set_cover_field_line(para, "校外导师", metadata.cosupervisor_cn)
        elif "学 位 类 别" in text:
            _set_cover_field_line(para, "学位类别", metadata.degree_label)
        elif text.startswith(school) and "楷体" in text:
            _set_cover_plain_text(para, school)
        elif "20XX年X月" in text:
            _set_cover_plain_text(para, metadata.submit_date_cn)

    _apply_english_cover_paragraphs(doc, metadata)

    table_value_map = {
        metadata.student_label_cn: metadata.author_cn,
        "指导教师": supervisor_cn,
        "校外导师": metadata.cosupervisor_cn,
        "专业类别": metadata.subject_cn,
        "所在学院": metadata.affiliation_cn,
        "论文提交日期": metadata.submit_date_cn,
        "论文答辩日期": metadata.oral_date_cn,
        "学位授予单位": school,
        **english_table_values,
    }
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = _cover_table_label(cells[0].text)
            if label in {"硕士研究生", "博士研究生"}:
                _set_cell_text(cells[0], metadata.student_label_cn)
                label = metadata.student_label_cn
            if label in table_value_map:
                value = table_value_map[label]
                prefix = "：" if "：" not in label and not label.endswith(":") else ""
                _set_cell_text(cells[1], f"{prefix}{value}")


def _set_declaration_two_column_line(paragraph, left_text: str, right_text: str) -> None:
    _clear_paragraph(paragraph)
    ppr = paragraph._p.get_or_add_pPr()
    for tag in (
        "w:numPr",
        "w:tabs",
        "w:ind",
        "w:keepNext",
        "w:keepLines",
        "w:pageBreakBefore",
    ):
        _remove_child(ppr, tag)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)

    left = paragraph.add_run(left_text)
    tab = paragraph.add_run()
    tab.add_tab()
    right = paragraph.add_run(right_text)
    for run in (left, tab, right):
        _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, False)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.LEFT)


def _format_declaration_template(doc: Document) -> None:
    date_text = "日期：      年   月   日"
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("作者(签字)") and "导师(签字)" in text:
            _set_declaration_two_column_line(para, "作者(签字)：", "导师(签字)：")
        elif text.startswith("日期") and text.count("日期") >= 2:
            _set_declaration_two_column_line(para, date_text, date_text)


def _prepare_template_doc(path: Path, metadata: ThesisMetadata, *, kind: str) -> Document | None:
    if not path.exists():
        return None
    template = Document(path)
    if kind == "cover":
        _apply_cover_metadata(template, metadata)
    elif kind == "declaration":
        _format_declaration_template(template)
    return template


def _copy_missing_template_styles(target: Document, template: Document) -> None:
    target_names = {style.name for style in target.styles}
    target_ids = {style.style_id for style in target.styles}
    styles_element = target.styles.element
    for style in template.styles:
        if style.name in target_names or style.style_id in target_ids:
            continue
        styles_element.append(deepcopy(style._element))
        target_names.add(style.name)
        target_ids.add(style.style_id)
    _format_cover_template_styles(target)


def _format_cover_template_styles(doc: Document) -> None:
    styles = doc.styles
    if "封面宋体小2" in styles:
        _set_style_font(styles["封面宋体小2"], "Times New Roman", "Times New Roman", 18.0, False)
    if "封面2号黑体英文论文名称" in styles:
        _set_style_font(
            styles["封面2号黑体英文论文名称"],
            "Times New Roman",
            "Times New Roman",
            22.0,
            False,
        )


def _template_body_elements(template: Document, *, boundary_type: str | None = "nextPage") -> list:
    body = template._element.body
    elements = []
    for child in body:
        if child.tag == qn("w:sectPr"):
            continue
        copied = deepcopy(child)
        _strip_header_footer_references(copied)
        elements.append(copied)
    if boundary_type == "page":
        elements.append(_page_break_paragraph())
    elif boundary_type:
        final_break = _section_break_paragraph(template, boundary_type)
        elements.append(final_break)
    return elements


def _insert_elements_before_paragraph(paragraph, elements: list) -> None:
    parent = paragraph._p.getparent()
    index = parent.index(paragraph._p)
    for element in elements:
        parent.insert(index, element)
        index += 1


def _is_followed_by_section_break(paragraph) -> bool:
    parent = paragraph._p.getparent()
    index = parent.index(paragraph._p)
    for following in parent[index + 1 :]:
        if _is_empty_section_break_paragraph(following):
            return True
        if following.tag == qn("w:p") and not _element_text(following).strip():
            continue
        return False
    return False


def _replace_marker_with_template(
    doc: Document,
    marker: str,
    template_path: Path,
    metadata: ThesisMetadata,
    report: ConversionReport,
    *,
    kind: str,
    note: str,
) -> bool:
    template = _prepare_template_doc(template_path, metadata, kind=kind)
    if template is None:
        report.warn(f"找不到 {note} 模板：{template_path}")
        return False

    for para in list(doc.paragraphs):
        if para.text.strip() == marker:
            _copy_missing_template_styles(doc, template)
            if marker == DECLARATION_MARKER:
                boundary_type = None
            else:
                boundary_type = None if _is_followed_by_section_break(para) else "page"
            _insert_elements_before_paragraph(para, _template_body_elements(template, boundary_type=boundary_type))
            _remove_paragraph(para)
            report.note(f"已插入{note}：{template_path}")
            return True
    return False


def _prepend_template(
    doc: Document,
    template_path: Path,
    metadata: ThesisMetadata,
    report: ConversionReport,
    *,
    kind: str,
    note: str,
) -> None:
    template = _prepare_template_doc(template_path, metadata, kind=kind)
    if template is None:
        report.warn(f"找不到 {note} 模板：{template_path}")
        return
    _copy_missing_template_styles(doc, template)
    body = doc._element.body
    insert_at = 0
    for element in _template_body_elements(template, boundary_type="page"):
        body.insert(insert_at, element)
        insert_at += 1
    report.note(f"已插入{note}：{template_path}")


def _style(doc: Document, name: str, base: str = "Normal"):
    styles = doc.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base in styles:
        style.base_style = styles[base]
    return style


def _toc_style(doc: Document, level: int):
    style_id = f"TOC{level}"
    style = next((item for item in doc.styles if item.style_id == style_id), None)
    if style is None:
        # python-docx gives this the reserved TOC styleId; removing
        # customStyle makes Word treat it as the corresponding built-in style
        # instead of renaming it and creating a second default TOC style.
        style = _style(doc, f"TOC {level}")
    style._element.attrib.pop(qn("w:customStyle"), None)
    name = style._element.find(qn("w:name"))
    if name is None:
        name = OxmlElement("w:name")
        style._element.insert(0, name)
    name.set(qn("w:val"), f"toc {level}")
    return style


def _show_style_in_gallery(style, priority: int) -> None:
    element = style._element
    _remove_child(element, "w:semiHidden")
    _remove_child(element, "w:unhideWhenUsed")
    _remove_child(element, "w:qFormat")
    _remove_child(element, "w:uiPriority")
    priority_el = OxmlElement("w:uiPriority")
    priority_el.set(qn("w:val"), str(priority))
    element.append(priority_el)
    element.append(OxmlElement("w:qFormat"))
    try:
        style.hidden = False
        style.quick_style = True
        style.priority = priority
    except AttributeError:
        pass


def _hide_style_from_gallery(style) -> None:
    element = style._element
    _remove_child(element, "w:qFormat")
    _remove_child(element, "w:uiPriority")
    if element.find(qn("w:semiHidden")) is None:
        element.append(OxmlElement("w:semiHidden"))
    if element.find(qn("w:unhideWhenUsed")) is None:
        element.append(OxmlElement("w:unhideWhenUsed"))
    try:
        style.hidden = True
        style.quick_style = False
    except AttributeError:
        pass


def _set_style_outline_level(style, level: int) -> None:
    ppr = style._element.get_or_add_pPr()
    _remove_child(ppr, "w:outlineLvl")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    ppr.append(outline)


def _set_paragraph_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    _remove_child(ppr, "w:outlineLvl")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    ppr.append(outline)


def _enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    _remove_child(settings, "w:updateFields")
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _referenced_style_ids(doc: Document) -> set[str]:
    style_ids: set[str] = set()
    for element in doc.element.iter():
        tag = element.tag
        if tag in {qn("w:pStyle"), qn("w:rStyle"), qn("w:tblStyle")}:
            val = element.get(qn("w:val"))
            if val:
                style_ids.add(val)
    return style_ids


def _remove_unrelated_custom_styles(doc: Document) -> None:
    referenced_ids = _referenced_style_ids(doc)
    styles_element = doc.styles.element
    for style in list(doc.styles):
        element = style._element
        if element.get(qn("w:customStyle")) != "1":
            continue
        if style.name in KEEP_CUSTOM_STYLE_NAMES:
            continue
        if style.style_id in referenced_ids:
            _hide_style_from_gallery(style)
            continue
        styles_element.remove(element)


def _normalize_pandoc_custom_styles(doc: Document) -> None:
    paragraph_styles = {"Compact", "Definition Term", "Definition"}
    run_styles = {"Verbatim Char"}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name in paragraph_styles:
            _set_paragraph_style(para, HEU_STYLES["body"])
        for run in para.runs:
            if run.style and run.style.name in run_styles:
                run.style = doc.styles["Default Paragraph Font"]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style_name = para.style.name if para.style else ""
                    if style_name in paragraph_styles:
                        _set_paragraph_style(para, HEU_STYLES["body"])
                    for run in para.runs:
                        if run.style and run.style.name in run_styles:
                            run.style = doc.styles["Default Paragraph Font"]


def _set_paragraph_style(para, style_name: str) -> None:
    try:
        para.style = style_name
    except KeyError:
        pass


def _bookmark_name(key: str) -> str:
    safe = re.sub(r"[^0-9A-Za-z_]", "_", key)
    if not safe or safe[0].isdigit():
        safe = f"ref_{safe}"
    return f"HEU_REF_{safe[:32]}"


def _next_bookmark_id(doc: Document) -> int:
    ids = []
    for start in doc._element.iter(qn("w:bookmarkStart")):
        value = start.get(qn("w:id"))
        if value is not None:
            try:
                ids.append(int(value))
            except ValueError:
                continue
    return max(ids, default=0) + 1


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def _create_reference_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(value)
        for element in numbering.findall(qn("w:abstractNum"))
        if (value := element.get(qn("w:abstractNumId"))) is not None and value.isdigit()
    ]
    num_ids = [
        int(value)
        for element in numbering.findall(qn("w:num"))
        if (value := element.get(qn("w:numId"))) is not None and value.isdigit()
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("w:start", "1"),
        ("w:numFmt", "decimal"),
        ("w:lvlText", "[%1]"),
        ("w:suff", "space"),
        ("w:lvlJc", "left"),
    ):
        child = OxmlElement(tag)
        child.set(qn("w:val"), value)
        level.append(child)
    abstract.append(level)
    numbering.append(abstract)

    instance = OxmlElement("w:num")
    instance.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    instance.append(abstract_ref)
    numbering.append(instance)
    return num_id


def _set_reference_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    _remove_child(ppr, "w:numPr")
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_pr.append(level)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(number)
    ppr.append(num_pr)


def _insert_paragraph_after(paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_para = OxmlElement("w:p")
    parent = paragraph._p.getparent()
    parent.insert(parent.index(paragraph._p) + 1, new_para)
    inserted = Paragraph(new_para, paragraph._parent)
    if style:
        _set_paragraph_style(inserted, style)
    if text:
        inserted.add_run(text)
    return inserted


def _insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(16.0))
    body = paragraph._p.getparent()
    table_parent = table._tbl.getparent()
    table_parent.remove(table._tbl)
    body.insert(body.index(paragraph._p) + 1, table._tbl)
    return table


def _insert_paragraph_after_table(table, parent, text: str = "", style: str | None = None) -> Paragraph:
    new_para = OxmlElement("w:p")
    body = table._tbl.getparent()
    body.insert(body.index(table._tbl) + 1, new_para)
    inserted = Paragraph(new_para, parent)
    if style:
        _set_paragraph_style(inserted, style)
    if text:
        inserted.add_run(text)
    return inserted


def _decode_citation_token(token: str) -> list[str]:
    keys: list[str] = []
    for part in token.split("."):
        padded = part + "=" * (-len(part) % 4)
        try:
            keys.append(base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8"))
        except Exception:
            continue
    return keys


def _decode_marker_token(token: str) -> str | None:
    padded = token + "=" * (-len(token) % 4)
    try:
        return base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8")
    except Exception:
        return None


def _equation_bookmark_name(label: str) -> str:
    safe = re.sub(r"[^0-9A-Za-z_]", "_", label)
    if not safe or safe[0].isdigit():
        safe = f"eq_{safe}"
    return f"HEU_EQ_{safe[:32]}"


def _format_sections(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(2.0)
        section.footer_distance = Cm(2.0)


def _format_styles(doc: Document) -> None:
    styles = doc.styles
    _style(doc, HEU_STYLES["body"])
    _style(doc, HEU_STYLES["body_inline_math"], HEU_STYLES["body"])
    _style(doc, HEU_STYLES["chapter"], "Heading 1")
    _style(doc, HEU_STYLES["section"], "Heading 2")
    _style(doc, HEU_STYLES["subsection"], "Heading 3")
    _style(doc, HEU_STYLES["subsubsection"], "Heading 4")
    _style(doc, HEU_STYLES["caption"], "Caption")
    _style(doc, HEU_STYLES["figure"])
    _style(doc, HEU_STYLES["display_math"])
    _style(doc, HEU_STYLES["reference"])
    _style(doc, HEU_STYLES["cover_title"])
    _style(doc, HEU_STYLES["cover_meta"])
    _style(doc, HEU_STYLES["cover_small"])
    toc_styles = [_toc_style(doc, level) for level in range(1, 4)]
    for legacy_name in LEGACY_HEU_STYLE_NAMES:
        if legacy_name in styles:
            _hide_style_from_gallery(styles[legacy_name])

    gallery_priorities = {
        HEU_STYLES["chapter"]: 1,
        HEU_STYLES["section"]: 2,
        HEU_STYLES["subsection"]: 3,
        HEU_STYLES["subsubsection"]: 4,
        HEU_STYLES["body"]: 10,
        HEU_STYLES["body_inline_math"]: 11,
        HEU_STYLES["caption"]: 20,
        HEU_STYLES["figure"]: 21,
        HEU_STYLES["display_math"]: 22,
        HEU_STYLES["reference"]: 30,
        HEU_STYLES["cover_title"]: 40,
        HEU_STYLES["cover_meta"]: 41,
        HEU_STYLES["cover_small"]: 42,
    }
    for style_name, priority in gallery_priorities.items():
        _show_style_in_gallery(styles[style_name], priority)

    for style_name, outline_level in (
        (HEU_STYLES["chapter"], 0),
        (HEU_STYLES["section"], 1),
        (HEU_STYLES["subsection"], 2),
        (HEU_STYLES["subsubsection"], 3),
    ):
        _set_style_outline_level(styles[style_name], outline_level)
        styles[style_name].paragraph_format.keep_with_next = True
        styles[style_name].paragraph_format.keep_together = True

    for name in ("Normal", "Body Text", "First Paragraph"):
        if name in styles:
            _set_style_font(styles[name], "宋体", "Times New Roman", BODY_SIZE_PT, False)
            _set_style_paragraph_layout(
                styles[name],
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent_pt=FIRST_LINE_INDENT_PT,
                line_spacing_pt=BODY_LINE_SPACING_PT,
                space_before_pt=0,
                space_after_pt=0,
            )
    _set_style_font(styles[HEU_STYLES["body"]], "宋体", "Times New Roman", BODY_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["body"]],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=FIRST_LINE_INDENT_PT,
        line_spacing_pt=BODY_LINE_SPACING_PT,
        space_before_pt=0,
        space_after_pt=0,
    )
    _set_style_font(styles[HEU_STYLES["body_inline_math"]], "宋体", "Times New Roman", BODY_SIZE_PT, False)
    _set_style_single_line_layout(
        styles[HEU_STYLES["body_inline_math"]],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=FIRST_LINE_INDENT_PT,
        space_before_pt=0,
        space_after_pt=0,
    )

    for name in ("Heading 1", "Title"):
        if name in styles:
            _set_style_font(styles[name], "黑体", "Times New Roman", CHAPTER_SIZE_PT, False)
            _set_style_paragraph_layout(
                styles[name],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_pt=0,
                line_spacing_pt=CHAPTER_LINE_SPACING_PT,
                space_before_pt=CHAPTER_SPACE_BEFORE_PT,
                space_after_pt=CHAPTER_SPACE_AFTER_PT,
            )
    _set_style_font(styles[HEU_STYLES["chapter"]], "黑体", "Times New Roman", CHAPTER_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["chapter"]],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        line_spacing_pt=CHAPTER_LINE_SPACING_PT,
        space_before_pt=CHAPTER_SPACE_BEFORE_PT,
        space_after_pt=CHAPTER_SPACE_AFTER_PT,
    )
    if "Heading 2" in styles:
        _set_style_font(styles["Heading 2"], "黑体", "Times New Roman", SECTION_SIZE_PT, False)
        _set_style_paragraph_layout(
            styles["Heading 2"],
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_pt=0,
            line_spacing_pt=SECTION_LINE_SPACING_PT,
            space_before_pt=SECTION_SPACE_PT,
            space_after_pt=SECTION_SPACE_PT,
        )
    _set_style_font(styles[HEU_STYLES["section"]], "黑体", "Times New Roman", SECTION_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["section"]],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        line_spacing_pt=SECTION_LINE_SPACING_PT,
        space_before_pt=SECTION_SPACE_PT,
        space_after_pt=SECTION_SPACE_PT,
    )
    if "Heading 3" in styles:
        _set_style_font(styles["Heading 3"], "黑体", "Times New Roman", SUBSECTION_SIZE_PT, False)
        _set_style_paragraph_layout(
            styles["Heading 3"],
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_pt=0,
            line_spacing_pt=SUBSECTION_LINE_SPACING_PT,
            space_before_pt=SUBSECTION_SPACE_PT,
            space_after_pt=SUBSECTION_SPACE_PT,
        )
    _set_style_font(styles[HEU_STYLES["subsection"]], "黑体", "Times New Roman", SUBSECTION_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["subsection"]],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        line_spacing_pt=SUBSECTION_LINE_SPACING_PT,
        space_before_pt=SUBSECTION_SPACE_PT,
        space_after_pt=SUBSECTION_SPACE_PT,
    )
    if "Heading 4" in styles:
        _set_style_font(styles["Heading 4"], "黑体", "Times New Roman", SUBSUBSECTION_SIZE_PT, False)
        _set_style_paragraph_layout(
            styles["Heading 4"],
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_pt=0,
            line_spacing_pt=SUBSUBSECTION_LINE_SPACING_PT,
            space_before_pt=SUBSUBSECTION_SPACE_PT,
            space_after_pt=SUBSUBSECTION_SPACE_PT,
        )
    _set_style_font(styles[HEU_STYLES["subsubsection"]], "黑体", "Times New Roman", SUBSUBSECTION_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["subsubsection"]],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        line_spacing_pt=SUBSUBSECTION_LINE_SPACING_PT,
        space_before_pt=SUBSUBSECTION_SPACE_PT,
        space_after_pt=SUBSUBSECTION_SPACE_PT,
    )

    for name in ("Caption", "Image Caption", "Table Caption"):
        if name in styles:
            _set_style_font(styles[name], "宋体", "Times New Roman", CAPTION_SIZE_PT, False)
            _set_style_paragraph_layout(
                styles[name],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_pt=0,
                line_spacing_pt=CAPTION_LINE_SPACING_PT,
                space_before_pt=0,
                space_after_pt=0,
            )
    _set_style_font(styles[HEU_STYLES["caption"]], "宋体", "Times New Roman", CAPTION_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["caption"]],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        line_spacing_pt=CAPTION_LINE_SPACING_PT,
        space_before_pt=0,
        space_after_pt=0,
    )

    toc1, toc2, toc3 = toc_styles
    _set_style_font(toc1, "黑体", "Times New Roman", BODY_SIZE_PT, False)
    _set_style_single_line_layout(
        toc1,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        space_before_pt=6,
        space_after_pt=6,
    )
    toc1.paragraph_format.left_indent = Pt(0)
    toc1.paragraph_format.right_indent = Pt(0)

    for style, left_indent_pt in ((toc2, 24), (toc3, 48)):
        _set_style_font(style, "宋体", "Times New Roman", BODY_SIZE_PT, False)
        _set_style_paragraph_layout(
            style,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=0,
            line_spacing_pt=22,
            space_before_pt=0,
            space_after_pt=0,
        )
        style.paragraph_format.left_indent = Pt(left_indent_pt)
        style.paragraph_format.right_indent = Pt(0)

    _set_style_font(styles[HEU_STYLES["figure"]], "宋体", "Times New Roman", BODY_SIZE_PT, False)
    _set_style_exact_line_layout(
        styles[HEU_STYLES["figure"]],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        line_spacing_pt=1.0,
        space_before_pt=0,
        space_after_pt=0,
    )
    _set_style_font(styles[HEU_STYLES["display_math"]], "宋体", "Times New Roman", BODY_SIZE_PT, False)
    _set_style_single_line_layout(
        styles[HEU_STYLES["display_math"]],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=8,
        space_after_pt=8,
    )
    _set_style_font(styles[HEU_STYLES["reference"]], "宋体", "Times New Roman", BODY_SIZE_PT, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["reference"]],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        line_spacing_pt=BODY_LINE_SPACING_PT,
        space_before_pt=0,
        space_after_pt=0,
    )
    _set_style_font(styles[HEU_STYLES["cover_title"]], "黑体", "Times New Roman", 22.0, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["cover_title"]],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        line_spacing_pt=28.0,
        space_before_pt=0,
        space_after_pt=0,
    )
    _set_style_font(styles[HEU_STYLES["cover_meta"]], "宋体", "Times New Roman", 14.0, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["cover_meta"]],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        line_spacing_pt=23.0,
        space_before_pt=0,
        space_after_pt=0,
    )
    _set_style_font(styles[HEU_STYLES["cover_small"]], "宋体", "Times New Roman", 12.0, False)
    _set_style_paragraph_layout(
        styles[HEU_STYLES["cover_small"]],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        line_spacing_pt=18.0,
        space_before_pt=0,
        space_after_pt=0,
    )


def _center_header_footer_paragraph(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    _remove_child(ppr, "w:tabs")
    _remove_child(ppr, "w:ind")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _set_header_decoration_line(paragraph) -> None:
    """Match the HEU header's compound double paragraph border."""
    ppr = paragraph._p.get_or_add_pPr()
    _remove_child(ppr, "w:pBdr")

    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "thickThinSmallGap" if edge == "bottom" else "none")
        border.set(qn("w:sz"), "24" if edge == "bottom" else "0")
        border.set(qn("w:space"), "1" if edge == "bottom" else "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)

    ppr.insert_element_before(
        borders,
        "w:shd",
        "w:tabs",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:outlineLvl",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )


def _format_headers_and_footers(doc: Document, metadata: ThesisMetadata) -> None:
    school_line = f"哈尔滨工程大学{metadata.degree_label}学位论文"
    section_titles = [
        _first_meaningful_section_text(blocks)
        for blocks, _sect_pr in _section_blocks(doc)
    ]
    main_matter_started = False
    for idx, section in enumerate(doc.sections):
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType")) if sect_pr is not None else None
        has_page_numbering = pg_num is not None
        for part in (
            section.header,
            section.even_page_header,
            section.first_page_header,
            section.footer,
            section.even_page_footer,
            section.first_page_footer,
        ):
            if idx > 0:
                # A copied sectPr can explicitly reuse the same header part while
                # reporting it as unlinked. Relink then unlink to force a new part.
                part.is_linked_to_previous = True
            part.is_linked_to_previous = False
        if idx == 0 or not has_page_numbering:
            for header in (section.header, section.even_page_header, section.first_page_header):
                para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                _clear_paragraph(para)
                _remove_child(para._p.get_or_add_pPr(), "w:pBdr")
            for footer in (section.footer, section.even_page_footer, section.first_page_footer):
                para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                _clear_paragraph(para)
            continue
        section.different_first_page_header_footer = False

        chapter_title = (
            section_titles[idx].strip()
            if idx < len(section_titles) and section_titles[idx].strip()
            else "章节标题"
        )
        if CHAPTER_PREFIX_RE.match(chapter_title):
            main_matter_started = True

        odd_header = section.header
        para = odd_header.paragraphs[0] if odd_header.paragraphs else odd_header.add_paragraph()
        _clear_paragraph(para)
        _center_header_footer_paragraph(para)
        if main_matter_started:
            _add_field(para, 'STYLEREF "章标题" \\* MERGEFORMAT', chapter_title)
        else:
            para.add_run(school_line)
        for run in para.runs:
            _set_run_font(run, "宋体", "Times New Roman", 10.5, False)
        _set_header_decoration_line(para)

        even_header = section.even_page_header
        para = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
        _clear_paragraph(para)
        _center_header_footer_paragraph(para)
        run = para.add_run(school_line)
        _set_run_font(run, "宋体", "Times New Roman", 10.5, False)
        _set_header_decoration_line(para)

        for footer in (section.footer, section.even_page_footer):
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            _clear_paragraph(para)
            _center_header_footer_paragraph(para)
            _add_field(para, "PAGE")
            for run in para.runs:
                _set_run_font(run, "宋体", "Times New Roman", 10.5, False)


def _add_center_paragraph_after(anchor: Paragraph, text: str, style: str, size: float, font_cn: str = "宋体", bold: bool = False) -> Paragraph:
    para = _insert_paragraph_after(anchor, text, style)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        _set_run_font(run, font_cn, "Times New Roman", size, bold)
    return para


def _fill_cover_table(table, rows: list[tuple[str, str]], *, label_width_cm: float = 4.0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for old in list(tbl_pr.findall(qn("w:tblBorders"))):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set_border(borders, edge, val="nil")
    tbl_pr.append(borders)
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].width = Cm(label_width_cm)
        row.cells[1].width = Cm(8.5)
        row.cells[0].text = label
        row.cells[1].text = value or " "
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in list(tc_pr.findall(qn("w:tcBorders"))):
                tc_pr.remove(old)
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                _set_border(tc_borders, edge, val="nil")
            tc_pr.append(tc_borders)
            for para in cell.paragraphs:
                _set_paragraph_style(para, HEU_STYLES["cover_meta"])
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_paragraph_layout(
                    para,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent_pt=0,
                    line_spacing_pt=23.0,
                    space_before_pt=0,
                    space_after_pt=0,
                )
                for run in para.runs:
                    _set_run_font(run, "宋体", "Times New Roman", 14.0, False)


def _fill_cover_bold_table(table, rows: list[tuple[str, str]], *, label_width_cm: float = 4.2) -> None:
    _fill_cover_table(table, rows, label_width_cm=label_width_cm)
    for row in table.rows:
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True


def _add_page_break_after(anchor: Paragraph) -> Paragraph:
    anchor.add_run().add_break(WD_BREAK.PAGE)
    return anchor


def _add_section_break_after(anchor: Paragraph) -> Paragraph:
    para = _insert_paragraph_after(anchor)
    ppr = para._p.get_or_add_pPr()
    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)
    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "11906")
    pg_sz.set(qn("w:h"), "16838")
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement("w:pgMar")
    for key, value in {
        "top": "1587",
        "right": "1417",
        "bottom": "1587",
        "left": "1417",
        "header": "1134",
        "footer": "1134",
        "gutter": "0",
    }.items():
        pg_mar.set(qn(f"w:{key}"), value)
    sect_pr.append(pg_mar)
    ppr.append(sect_pr)
    return para


def _add_cover_top_info(anchor: Paragraph, metadata: ThesisMetadata) -> Paragraph:
    def add_line(after: Paragraph, pairs: tuple[tuple[str, str], tuple[str, str]]) -> Paragraph:
        para = _insert_paragraph_after(after, "", HEU_STYLES["cover_small"])
        para.paragraph_format.left_indent = Cm(1.0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for idx, (label, value) in enumerate(pairs):
            label_run = para.add_run(label)
            _set_run_font(label_run, "宋体", "Times New Roman", 12.0, False)
            value_run = para.add_run(f"  {value or ' '}  ")
            value_run.font.underline = True
            _set_run_font(value_run, "宋体", "Times New Roman", 12.0, False)
            if idx == 0:
                gap = para.add_run(" " * 32)
                _set_run_font(gap, "宋体", "Times New Roman", 12.0, False)
        _set_paragraph_layout(
            para,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_pt=0,
            line_spacing_pt=15.0,
            space_before_pt=0,
            space_after_pt=0,
        )
        return para

    line1 = add_line(anchor, (("分类号：", metadata.classified_index), ("密级：", metadata.secret_level)))
    return add_line(line1, (("U D C：", metadata.udc), ("编号：", metadata.document_number)))


def _build_chinese_cover(anchor: Paragraph, metadata: ThesisMetadata) -> Paragraph:
    current = anchor
    _set_paragraph_style(current, HEU_STYLES["cover_small"])
    current.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current.paragraph_format.space_before = Pt(26)
    current = _add_cover_top_info(current, metadata)
    cover_label = "工学博士学位论文" if metadata.degree == "doctor" else "工学硕士学位论文"
    current = _add_center_paragraph_after(current, cover_label, HEU_STYLES["cover_meta"], 18.0)
    current.paragraph_format.space_before = Pt(54)
    current = _add_center_paragraph_after(current, metadata.display_title_cn, HEU_STYLES["cover_title"], 22.0, "黑体", True)
    current.paragraph_format.space_before = Pt(45)
    current.paragraph_format.space_after = Pt(125)
    rows = [
        ("硕 士 研 究 生：", metadata.author_cn),
        ("指  导  教  师：", metadata.supervisor_cn),
        ("副    导    师：", metadata.associate_supervisor_cn or metadata.co_supervisor_cn or ""),
    ]
    rows.extend(
        [
            ("学  科、专  业：", metadata.subject_cn),
            ("学位论文主审人：", metadata.co_supervisor_cn or metadata.supervisor_cn),
        ]
    )
    table = _insert_table_after(current, len(rows), 2)
    _fill_cover_table(table, rows, label_width_cm=4.4)
    current = _insert_paragraph_after_table(table, current._parent, "", HEU_STYLES["cover_meta"])
    current.paragraph_format.space_before = Pt(128)
    current.add_run("哈尔滨工程大学")
    for run in current.runs:
        _set_run_font(run, "楷体", "Times New Roman", 18.0, False)
    current.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current = _add_center_paragraph_after(current, metadata.submit_date_cn or "2026 年 7 月", HEU_STYLES["cover_meta"], 15.0)
    return _add_page_break_after(current)


def _build_blank_cover_page(anchor: Paragraph) -> Paragraph:
    current = _insert_paragraph_after(anchor, "", HEU_STYLES["cover_small"])
    current.paragraph_format.space_before = Pt(1)
    current.add_run().add_break(WD_BREAK.PAGE)
    return current


def _build_chinese_inner_cover(anchor: Paragraph, metadata: ThesisMetadata) -> Paragraph:
    current = _insert_paragraph_after(anchor, "", HEU_STYLES["cover_small"])
    current.paragraph_format.space_before = Pt(0)
    current = _add_cover_top_info(current, metadata)
    cover_label = "工学博士学位论文" if metadata.degree == "doctor" else "工学硕士学位论文"
    current = _add_center_paragraph_after(current, cover_label, HEU_STYLES["cover_meta"], 18.0)
    current.paragraph_format.space_before = Pt(70)
    current = _add_center_paragraph_after(current, metadata.display_title_cn, HEU_STYLES["cover_title"], 22.0, "黑体", True)
    current.paragraph_format.space_before = Pt(46)
    current.paragraph_format.space_after = Pt(97)
    rows = [
        ("硕 士 研 究 生：", metadata.author_cn),
        ("指  导  教  师：", metadata.supervisor_cn),
        ("副    导    师：", metadata.associate_supervisor_cn or metadata.co_supervisor_cn or ""),
        ("申  请  学  位：", f"{metadata.discipline_cn or '工学'}{'博士' if metadata.degree == 'doctor' else '硕士'}"),
        ("学  科、专  业：", metadata.subject_cn),
        ("所  在  单  位：", metadata.affiliation_cn),
        ("论文提交日期：", metadata.submit_date_cn),
        ("论文答辩日期：", metadata.oral_date_cn),
        ("学位授予单位：", "哈尔滨工程大学"),
    ]
    table = _insert_table_after(current, len(rows), 2)
    _fill_cover_bold_table(table, rows, label_width_cm=4.8)
    current = _insert_paragraph_after_table(table, current._parent)
    return _add_page_break_after(current)


def _build_english_cover(anchor: Paragraph, metadata: ThesisMetadata) -> Paragraph:
    current = _insert_paragraph_after(anchor, "", HEU_STYLES["cover_small"])
    current.add_run(f"Classified Index: {metadata.classified_index or ' '}        ")
    current.add_run(f"U.D.C: {metadata.udc or ' '}")
    degree = metadata.student_type_en or ("Doctor of Engineering" if metadata.degree == "doctor" else "Master of Engineering")
    degree_line = "A Dissertation for the Degree of D.Eng" if metadata.degree == "doctor" else "A Dissertation for the Degree of M.Eng"
    current = _add_center_paragraph_after(current, degree_line, HEU_STYLES["cover_meta"], 18.0)
    current.paragraph_format.space_before = Pt(92)
    current = _add_center_paragraph_after(current, metadata.title_en or metadata.display_title_cn, HEU_STYLES["cover_title"], 22.0, "Times New Roman")
    current.paragraph_format.space_before = Pt(42)
    current.paragraph_format.space_after = Pt(150)
    associate_supervisor = metadata.associate_supervisor_en or metadata.co_supervisor_en
    rows = [
        ("Candidate:", metadata.author_en or metadata.author_cn),
        ("Supervisor:", metadata.supervisor_en or metadata.supervisor_cn),
        *([("Associate Supervisor:", associate_supervisor)] if associate_supervisor else []),
        ("Academic Degree Applied for:", degree),
        ("Specialty:", metadata.subject_en or metadata.subject_cn),
        ("Affiliation:", metadata.affiliation_en or metadata.affiliation_cn),
        ("Date of Submission:", metadata.submit_date_en or metadata.submit_date_cn),
        ("Date of Oral Examination:", metadata.oral_date_en or metadata.oral_date_cn),
        ("University:", "Harbin Engineering University"),
    ]
    table = _insert_table_after(current, len(rows), 2)
    _fill_cover_table(table, rows, label_width_cm=6.0)
    current = _insert_paragraph_after_table(table, current._parent)
    return _add_section_break_after(current)


def _replace_cover_placeholder(doc: Document, metadata: ThesisMetadata, report: ConversionReport) -> None:
    for para in doc.paragraphs:
        if para.text.strip() != HEU_COVER_PLACEHOLDER:
            continue
        _clear_paragraph(para)
        current = _build_chinese_cover(para, metadata)
        current = _build_blank_cover_page(current)
        current = _build_chinese_inner_cover(current, metadata)
        current = _build_blank_cover_page(current)
        _build_english_cover(current, metadata)
        report.note("已生成可编辑封面、中文内封和英文内封。")
        missing = [
            label
            for label, value in {
                "题名": metadata.display_title_cn,
                "作者": metadata.author_cn,
                "导师": metadata.supervisor_cn,
                "学科专业": metadata.subject_cn,
                "学院": metadata.affiliation_cn,
                "提交日期": metadata.submit_date_cn,
                "答辩日期": metadata.oral_date_cn,
            }.items()
            if not value
        ]
        if missing:
            report.warn("封面字段缺失: " + "、".join(missing))
        return


def _replace_placeholders(
    doc: Document,
    report: ConversionReport,
    references: list[ReferenceEntry] | None = None,
) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in {"newpage", "ewpage"}:
            _clear_paragraph(para)
            continue
        if text == "HEU_TOC_PLACEHOLDER":
            _clear_paragraph(para)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_field(para, r'TOC \o "1-3" \h \z \u', "请更新目录")
            report.note(
                "已插入 Word 目录域并设置打开时自动更新；"
                "若 Word/WPS 未刷新，请按 Ctrl+A、F9 更新全部域。"
            )
        elif text == "HEU_REFERENCES_PLACEHOLDER":
            _clear_paragraph(para)
            if references:
                num_id = _create_reference_numbering(doc)
                bookmark_id = _next_bookmark_id(doc)
                _populate_reference_paragraph(
                    para,
                    references[0],
                    num_id=num_id,
                    bookmark_id=bookmark_id,
                )
                parent = para._p.getparent()
                index = parent.index(para._p)
                for ref in references[1:]:
                    new_para = OxmlElement("w:p")
                    parent.insert(index + 1, new_para)
                    index += 1
                    inserted = Paragraph(new_para, para._parent)
                    bookmark_id += 1
                    _populate_reference_paragraph(
                        inserted,
                        ref,
                        num_id=num_id,
                        bookmark_id=bookmark_id,
                    )
                report.note(
                    f"已从 BibTeX 生成 {len(references)} 条 Word 自动编号参考文献，"
                    "可在交叉引用的‘编号项’中选择。"
                )
            else:
                run = para.add_run("参考文献由 Pandoc citeproc 生成；若此处为空，请检查 BibTeX 或转换报告。")
                _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, False)
                report.warn("参考文献占位仍存在，Pandoc 未能自动生成参考文献列表。")
        elif text.startswith("HEU_LIST_OF_"):
            _clear_paragraph(para)
            run = para.add_run("请在 Word 中根据题注更新清单。")
            _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, False)


def _is_front_matter_title(text: str) -> bool:
    return text.strip() in FRONT_MATTER_TITLES


def _is_unnumbered_primary_title(text: str) -> bool:
    stripped = text.strip()
    return stripped in FRONT_MATTER_TITLES or stripped in PRIMARY_BACK_MATTER_TITLES or stripped in PUBLICATION_SECTION_TITLES


def _is_caption_style(style_name: str, kind: str | None = None) -> bool:
    if kind == "figure":
        return style_name in {"Caption", "Image Caption", "Figure Caption"}
    if kind == "table":
        return style_name == "Table Caption"
    return style_name in {"Caption", "Image Caption", "Figure Caption", "Table Caption"}


def _paragraph_has_drawing(para) -> bool:
    xml = para._element.xml
    return "<w:drawing" in xml or "<v:imagedata" in xml or "<wp:inline" in xml


def _paragraph_has_display_math(para) -> bool:
    return "<m:oMathPara" in para._element.xml


def _paragraph_has_math(para) -> bool:
    return "<m:oMath" in para._element.xml


def _wp_child(element, tag: str):
    return element.find(qn(tag))


def _convert_inline_to_top_bottom_anchor(inline) -> None:
    anchor = OxmlElement("wp:anchor")
    for attr, value in {
        "distT": "0",
        "distB": "0",
        "distL": "114300",
        "distR": "114300",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "0",
    }.items():
        anchor.set(attr, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    align = OxmlElement("wp:align")
    align.text = "center"
    position_h.append(align)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_offset = OxmlElement("wp:posOffset")
    pos_offset.text = "0"
    position_v.append(pos_offset)
    anchor.append(position_v)

    for tag in ("wp:extent", "wp:effectExtent"):
        child = _wp_child(inline, tag)
        if child is not None:
            anchor.append(deepcopy(child))

    wrap = OxmlElement("wp:wrapTopAndBottom")
    wrap.set("distT", "0")
    wrap.set("distB", "0")
    anchor.append(wrap)

    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        child = _wp_child(inline, tag)
        if child is not None:
            anchor.append(deepcopy(child))

    inline.getparent().replace(inline, anchor)


def _convert_images_to_top_bottom_wrap(doc: Document, report: ConversionReport | None = None) -> None:
    inlines = list(doc._element.iter(qn("wp:inline")))
    for inline in inlines:
        _convert_inline_to_top_bottom_anchor(inline)
    if report and inlines:
        report.note(f"已将 {len(inlines)} 张图片设置为上下型环绕。")


def _scale_inline_drawings(para, max_width_cm: float = 14.5) -> None:
    max_cx = int(max_width_cm * 360000)
    for extent in para._element.findall(".//" + qn("wp:extent")):
        cx = int(extent.get("cx", "0"))
        cy = int(extent.get("cy", "0"))
        if cx > max_cx and cx > 0:
            extent.set("cx", str(max_cx))
            extent.set("cy", str(int(cy * max_cx / cx)))
    for ext in para._element.findall(".//" + qn("a:ext")):
        cx = int(ext.get("cx", "0"))
        cy = int(ext.get("cy", "0"))
        if cx > max_cx and cx > 0:
            ext.set("cx", str(max_cx))
            ext.set("cy", str(int(cy * max_cx / cx)))


def _replace_text(para, text: str) -> None:
    _clear_paragraph(para)
    para.add_run(text)


def _replace_citation_placeholders(
    doc: Document,
    references: list[ReferenceEntry] | None,
    report: ConversionReport,
) -> None:
    ref_by_key = {ref.key: ref for ref in references or []}
    unresolved: set[str] = set()
    replaced = 0
    for para in doc.paragraphs:
        text = para.text
        if "HEU_CITE_" not in text:
            continue
        parts: list[str | tuple[str, str]] = []
        pos = 0
        for match in CITATION_PLACEHOLDER_RE.finditer(text):
            if match.start() > pos:
                parts.append(text[pos : match.start()])
            kind, token = match.group(1), match.group(2)
            keys = _decode_citation_token(token)
            parts.append((kind, ",".join(keys)))
            pos = match.end()
        if pos < len(text):
            parts.append(text[pos:])
        if not parts:
            continue

        _clear_paragraph(para)
        for part in parts:
            if isinstance(part, str):
                if part:
                    new_run = para.add_run(part)
                    _set_run_font(new_run, "宋体", "Times New Roman", BODY_SIZE_PT, None)
                continue
            kind, key_blob = part
            keys = [key for key in key_blob.split(",") if key]
            found_by_index = {
                ref_by_key[key].index: ref_by_key[key]
                for key in keys
                if key in ref_by_key
            }
            found = [found_by_index[index] for index in sorted(found_by_index)]
            missing = [key for key in keys if key not in ref_by_key]
            unresolved.update(missing)
            if not found:
                new_run = para.add_run("[?]")
                _set_run_font(new_run, "宋体", "Times New Roman", BODY_SIZE_PT, None)
                if kind == "SUPER":
                    new_run.font.superscript = True
                continue
            citation_runs = [para.add_run("[")]
            for idx, ref in enumerate(found):
                if idx:
                    citation_runs.append(para.add_run(","))
                field_run = _add_ref_field(para, _bookmark_name(ref.key), str(ref.index))
                citation_runs.append(field_run)
            citation_runs.append(para.add_run("]"))
            for run in citation_runs:
                _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, None)
                if kind == "SUPER":
                    run.font.superscript = True
            replaced += 1
    if replaced:
        report.note(f"已将 {replaced} 处文献引用替换为 Word REF 交叉引用域。")
    if unresolved:
        report.warn("未解析的文献 citekey: " + "、".join(sorted(unresolved)))


def _format_reference_paragraph(para) -> None:
    _set_paragraph_style(para, HEU_STYLES["reference"])
    _clear_text_direct_formatting(para)


def _populate_reference_paragraph(
    paragraph,
    reference: ReferenceEntry,
    *,
    num_id: int,
    bookmark_id: int,
) -> None:
    _format_reference_paragraph(paragraph)
    _set_reference_numbering(paragraph, num_id)
    prefix = f"[{reference.index}]"
    body = reference.text[len(prefix) :].lstrip() if reference.text.startswith(prefix) else reference.text
    run = paragraph.add_run(body)
    _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, False)
    _add_bookmark(paragraph, _bookmark_name(reference.key), bookmark_id)


def _format_heading_paragraph(
    para,
    level: int,
    text: str,
    chapter_no: int | None = None,
    section_no: int | None = None,
    subsection_no: int | None = None,
    exclude_from_toc: bool = False,
) -> None:
    front_title = _is_front_matter_title(text)
    if level == 1 or front_title:
        if chapter_no is not None and not CHAPTER_PREFIX_RE.match(text):
            _replace_text(para, f"第{chapter_no}章 {text}")
        _set_paragraph_style(para, HEU_STYLES["chapter"])
    elif level == 2:
        if chapter_no is not None and section_no is not None:
            prefix = f"{chapter_no}.{section_no}"
            if not re.match(rf"^{re.escape(prefix)}(?:\s+|$)", text):
                _replace_text(para, f"{prefix} {text}")
        _set_paragraph_style(para, HEU_STYLES["section"])
    elif level == 3:
        if chapter_no is not None and section_no is not None and subsection_no is not None:
            prefix = f"{chapter_no}.{section_no}.{subsection_no}"
            if not re.match(rf"^{re.escape(prefix)}(?:\s+|$)", text):
                _replace_text(para, f"{prefix} {text}")
        _set_paragraph_style(para, HEU_STYLES["subsection"])
    else:
        _set_paragraph_style(para, HEU_STYLES["subsubsection"])
    _clear_text_direct_formatting(para)
    if exclude_from_toc or text.strip() == "目录":
        # Keep pre-main-matter titles visually chapter-like without listing them in the TOC.
        _set_paragraph_outline_level(para, 9)


def _format_caption_paragraph(para, *, kind: str, chapter_no: int, caption_no: int) -> None:
    text = para.text.strip()
    if text and not CAPTION_PREFIX_RE.match(text):
        prefix = "图" if kind == "figure" else "表"
        _replace_text(para, f"{prefix} {chapter_no}.{caption_no} {text}")

    _set_paragraph_style(para, HEU_STYLES["caption"])
    _clear_text_direct_formatting(para)


def _format_body_paragraph(para) -> None:
    _set_paragraph_style(para, HEU_STYLES["body"])
    _clear_text_direct_formatting(para)


def _format_paragraphs(doc: Document, report: ConversionReport | None = None) -> None:
    current_chapter = 0
    current_section = 0
    current_subsection = 0
    in_numbered_chapter = False
    figure_counts: dict[int, int] = {}
    table_counts: dict[int, int] = {}
    numbered_sections = 0
    numbered_subsections = 0
    numbered_figures = 0
    numbered_tables = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if style_name in {HEU_STYLES["cover_title"], HEU_STYLES["cover_meta"], HEU_STYLES["cover_small"]}:
            continue

        if style_name.startswith("Heading 1"):
            numbered_chapter = not _is_unnumbered_primary_title(text) or text == "参考文献"
            current_section = 0
            current_subsection = 0
            if numbered_chapter:
                current_chapter += 1
                in_numbered_chapter = True
                _format_heading_paragraph(para, 1, text, current_chapter)
            else:
                in_numbered_chapter = False
                _format_heading_paragraph(
                    para,
                    1,
                    text,
                    exclude_from_toc=current_chapter == 0,
                )
            continue
        if style_name.startswith("Heading 2"):
            if _is_front_matter_title(text):
                in_numbered_chapter = False
                _format_heading_paragraph(
                    para,
                    2,
                    text,
                    exclude_from_toc=current_chapter == 0,
                )
            elif in_numbered_chapter:
                current_section += 1
                current_subsection = 0
                numbered_sections += 1
                _format_heading_paragraph(
                    para,
                    2,
                    text,
                    chapter_no=current_chapter,
                    section_no=current_section,
                )
            else:
                _format_heading_paragraph(para, 2, text)
            continue
        if style_name.startswith("Heading 3"):
            if in_numbered_chapter and current_section:
                current_subsection += 1
                numbered_subsections += 1
                _format_heading_paragraph(
                    para,
                    3,
                    text,
                    chapter_no=current_chapter,
                    section_no=current_section,
                    subsection_no=current_subsection,
                )
            else:
                _format_heading_paragraph(para, 3, text)
            continue
        if style_name.startswith("Heading"):
            _format_heading_paragraph(para, 4, text)
            continue

        chapter_no = max(current_chapter, 1)
        if _is_caption_style(style_name, "figure"):
            figure_counts[chapter_no] = figure_counts.get(chapter_no, 0) + 1
            if not CAPTION_PREFIX_RE.match(text):
                numbered_figures += 1
            _format_caption_paragraph(
                para,
                kind="figure",
                chapter_no=chapter_no,
                caption_no=figure_counts[chapter_no],
            )
            continue
        if _is_caption_style(style_name, "table"):
            table_counts[chapter_no] = table_counts.get(chapter_no, 0) + 1
            if not CAPTION_PREFIX_RE.match(text):
                numbered_tables += 1
            _format_caption_paragraph(
                para,
                kind="table",
                chapter_no=chapter_no,
                caption_no=table_counts[chapter_no],
            )
            continue

        if style_name == HEU_STYLES["reference"] or REFERENCE_ENTRY_RE.match(text):
            _format_reference_paragraph(para)
            continue

        if _paragraph_has_drawing(para) and not text:
            _set_paragraph_style(para, HEU_STYLES["figure"])
            _clear_text_direct_formatting(para)
            _scale_inline_drawings(para)
            continue

        if _paragraph_has_display_math(para):
            _set_paragraph_style(para, HEU_STYLES["display_math"])
            _clear_text_direct_formatting(para)
            continue

        if _paragraph_has_math(para):
            _set_paragraph_style(para, HEU_STYLES["body_inline_math"])
            _clear_text_direct_formatting(para)
            continue

        _format_body_paragraph(para)

    if report and (numbered_figures or numbered_tables):
        report.note(f"已按章补全题注编号：图 {numbered_figures} 个，表 {numbered_tables} 个。")
    if report and (numbered_sections or numbered_subsections):
        report.note(
            f"已补全标题编号：节 {numbered_sections} 个，小节 {numbered_subsections} 个。"
        )


def _append_tab(paragraph) -> None:
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:tab"))


def _set_equation_tabs(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    _remove_child(ppr, "w:tabs")
    tabs = OxmlElement("w:tabs")
    for alignment, position in (("center", 4536), ("right", 9072)):
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), alignment)
        tab.set(qn("w:leader"), "none")
        tab.set(qn("w:pos"), str(position))
        tabs.append(tab)
    ppr.insert_element_before(
        tabs,
        "w:spacing",
        "w:ind",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:outlineLvl",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )


def _bookmark_equation_number(
    paragraph,
    start_run,
    end_run,
    labels: list[str],
    bookmark_id: int,
) -> int:
    for label in labels:
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), _equation_bookmark_name(label))
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        start_run._r.addprevious(start)
        end_run._r.addnext(end)
        bookmark_id += 1
    return bookmark_id


def _format_numbered_equation(
    paragraph,
    *,
    chapter_no: int,
    equation_no: int,
    labels: list[str],
    bookmark_id: int,
) -> int:
    math_para = paragraph._p.find(qn("m:oMathPara"))
    if math_para is None:
        return bookmark_id
    maths = [deepcopy(element) for element in math_para.findall(qn("m:oMath"))]
    paragraph._p.remove(math_para)
    _set_equation_tabs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _append_tab(paragraph)
    for math in maths:
        paragraph._p.append(math)
    _append_tab(paragraph)

    opening = paragraph.add_run("(")
    prefix = paragraph.add_run(f"{chapter_no}-")
    sequence = _add_field(
        paragraph,
        f" SEQ HEUEquation{chapter_no} \\* ARABIC ",
        str(equation_no),
    )
    suffix = paragraph.add_run(")")
    for run in (opening, prefix, sequence, suffix):
        _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, False)
    return _bookmark_equation_number(
        paragraph,
        prefix,
        sequence,
        labels,
        bookmark_id,
    )


def _number_equations(doc: Document, report: ConversionReport | None = None) -> dict[str, str]:
    paragraphs = list(doc.paragraphs)
    current_chapter = 0
    equation_counts: dict[int, int] = {}
    equation_numbers: dict[str, str] = {}
    bookmark_id = _next_bookmark_id(doc)
    numbered = 0
    orphan_markers = 0

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        chapter_match = CHAPTER_PREFIX_RE.match(text)
        if chapter_match:
            value = chapter_match.group(1)
            current_chapter = int(value) if value.isdigit() else current_chapter + 1
            continue
        if not _paragraph_has_display_math(paragraph):
            continue
        if index + 1 >= len(paragraphs):
            continue
        marker_paragraph = paragraphs[index + 1]
        marker_match = EQUATION_MARKER_RE.match(marker_paragraph.text.strip())
        if marker_match is None:
            continue

        chapter_no = max(current_chapter, 1)
        equation_no = equation_counts.get(chapter_no, 0) + 1
        equation_counts[chapter_no] = equation_no
        token_blob = marker_match.group(1)
        labels = [
            label
            for token in token_blob.split(".")
            if token != "NONE" and (label := _decode_marker_token(token))
        ]
        bookmark_id = _format_numbered_equation(
            paragraph,
            chapter_no=chapter_no,
            equation_no=equation_no,
            labels=labels,
            bookmark_id=bookmark_id,
        )
        for label in labels:
            equation_numbers[label] = f"{chapter_no}-{equation_no}"
        _remove_paragraph(marker_paragraph)
        numbered += 1

    for paragraph in list(doc.paragraphs):
        if EQUATION_MARKER_RE.match(paragraph.text.strip()):
            orphan_markers += 1
            _remove_paragraph(paragraph)

    if report and numbered:
        report.note(f"已添加按章公式编号 {numbered} 个。")
    if report and orphan_markers:
        report.warn(f"有 {orphan_markers} 个公式编号标记未匹配到可编辑 Word 公式。")
    return equation_numbers


def _replace_equation_reference_placeholders(
    doc: Document,
    equation_numbers: dict[str, str],
    report: ConversionReport | None = None,
) -> None:
    replaced = 0
    unresolved: set[str] = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "HEU_EQREF_" not in text:
            continue
        parts: list[str | tuple[str, str]] = []
        position = 0
        for match in EQUATION_REF_PLACEHOLDER_RE.finditer(text):
            if match.start() > position:
                parts.append(text[position : match.start()])
            kind, token = match.group(1), match.group(2)
            label = _decode_marker_token(token)
            if label is not None:
                parts.append((kind, label))
            position = match.end()
        if position < len(text):
            parts.append(text[position:])
        if not parts:
            continue

        _clear_paragraph(paragraph)
        for part in parts:
            if isinstance(part, str):
                if part:
                    run = paragraph.add_run(part)
                    _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, None)
                continue
            kind, label = part
            fallback = equation_numbers.get(label, "?")
            if label not in equation_numbers:
                unresolved.add(label)
            if kind == "PAREN":
                run = paragraph.add_run("(")
                _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, None)
            field = _add_field(
                paragraph,
                f" REF {_equation_bookmark_name(label)} \\h ",
                fallback,
            )
            _set_run_font(field, "宋体", "Times New Roman", BODY_SIZE_PT, None)
            if kind == "PAREN":
                run = paragraph.add_run(")")
                _set_run_font(run, "宋体", "Times New Roman", BODY_SIZE_PT, None)
            replaced += 1

    if report and replaced:
        report.note(f"已将 {replaced} 处公式引用替换为 Word REF 域。")
    if report and unresolved:
        report.warn("未解析的公式引用: " + "、".join(sorted(unresolved)))


def _keep_figures_with_captions(doc: Document, report: ConversionReport | None = None) -> None:
    linked = 0
    for caption in doc.paragraphs:
        style_name = caption.style.name if caption.style else ""
        if style_name != HEU_STYLES["caption"]:
            continue
        previous = caption._p.getprevious()
        while previous is not None and previous.tag not in {qn("w:p"), qn("w:tbl")}:
            previous = previous.getprevious()
        if previous is None:
            continue
        if previous.tag == qn("w:p") and previous.find(".//" + qn("w:drawing")) is not None:
            _set_keep_next(previous)
            _set_keep_lines(previous)
            linked += 1
        elif previous.tag == qn("w:tbl") and previous.find(".//" + qn("w:drawing")) is not None:
            rows = previous.findall(qn("w:tr"))
            if rows:
                for image_paragraph in rows[-1].findall(".//" + qn("w:p")):
                    _set_keep_next(image_paragraph)
                    _set_keep_lines(image_paragraph)
                linked += 1
        _set_keep_lines(caption._p)
    if report and linked:
        report.note(f"已绑定图片与题注同页 {linked} 组。")


def _has_meaningful_content_before(paragraph) -> bool:
    parent = paragraph._p.getparent()
    index = parent.index(paragraph._p)
    for previous in parent[:index]:
        if previous.tag == qn("w:tbl"):
            return True
        if previous.tag == qn("w:p") and (
            _element_text(previous).strip()
            or previous.find(".//" + qn("w:drawing")) is not None
            or previous.find(".//" + qn("w:br")) is not None
        ):
            return True
    return False


def _is_primary_page_heading(para) -> bool:
    text = para.text.strip()
    style_name = para.style.name if para.style else ""
    return (
        style_name.startswith("Heading 1")
        or style_name == HEU_STYLES["chapter"]
        or CHAPTER_PREFIX_RE.match(text) is not None
        or text in FRONT_MATTER_TITLES
        or text in PRIMARY_BACK_MATTER_TITLES
    )


def _insert_page_section_breaks(doc: Document, report: ConversionReport | None = None) -> None:
    odd_breaks = 0
    normal_breaks = 0
    roman_started = False
    decimal_started = False

    for para in list(doc.paragraphs):
        text = para.text.strip()
        if not _is_primary_page_heading(para):
            continue
        if not _has_meaningful_content_before(para) or _has_previous_section_break(para):
            sect_pr = _previous_section_properties(para)
        else:
            if text in ROMAN_PAGE_TITLES:
                sect_pr = _ensure_section_break_before(para, "nextPage")
                normal_breaks += 1
            else:
                sect_pr = _ensure_section_break_before(para, "oddPage")
                odd_breaks += 1

        if sect_pr is None:
            continue
        if text in ROMAN_PAGE_TITLES:
            if text == "摘要" and not roman_started:
                _set_section_page_numbering(sect_pr, "upperRoman", 1)
                roman_started = True
            elif roman_started:
                _set_section_page_numbering(sect_pr, "upperRoman", None)
        elif CHAPTER_PREFIX_RE.match(text) and not decimal_started:
            _set_section_page_numbering(sect_pr, "decimal", 1)
            decimal_started = True
        elif decimal_started:
            _set_section_page_numbering(sect_pr, "decimal", None)

    if report and (odd_breaks or normal_breaks):
        report.note(f"已插入分页控制：章标题奇数页 {odd_breaks} 处，前置标题新页 {normal_breaks} 处，并设置前置罗马页码与正文阿拉伯页码。")


def _set_border(parent, edge: str, *, val: str, size_eighths_pt: int = 0) -> None:
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), str(size_eighths_pt))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "auto")
    parent.append(border)


def _format_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    for old in list(tbl_pr.findall(qn("w:tblBorders"))):
        tbl_pr.remove(old)
    tbl_borders = OxmlElement("w:tblBorders")
    _set_border(tbl_borders, "top", val="single", size_eighths_pt=12)
    _set_border(tbl_borders, "bottom", val="single", size_eighths_pt=12)
    for edge in ("left", "right", "insideV", "insideH"):
        _set_border(tbl_borders, edge, val="nil")
    tbl_pr.append(tbl_borders)

    if not table.rows:
        return
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in list(tc_pr.findall(qn("w:tcBorders"))):
            tc_pr.remove(old)
        tc_borders = OxmlElement("w:tcBorders")
        _set_border(tc_borders, "bottom", val="single", size_eighths_pt=8)
        tc_pr.append(tc_borders)


def _clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    for old in list(tbl_pr.findall(qn("w:tblBorders"))):
        tbl_pr.remove(old)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideV", "insideH"):
        _set_border(tbl_borders, edge, val="nil")
    tbl_pr.append(tbl_borders)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in list(tc_pr.findall(qn("w:tcBorders"))):
                tc_pr.remove(old)
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "bottom", "left", "right", "insideV", "insideH"):
                _set_border(tc_borders, edge, val="nil")
            tc_pr.append(tc_borders)


def _format_tables(doc: Document) -> None:
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "<w:drawing" in table._tbl.xml or "<wp:inline" in table._tbl.xml:
            _clear_table_borders(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _set_paragraph_style(para, HEU_STYLES["figure"])
                        _clear_text_direct_formatting(para)
            continue
        if any(
            token in table_text
            for token in (
                "分类",
                "密级",
                "U D C",
                "硕 士 研 究 生",
                "论文提交日期",
                "Candidate:",
                "Classified Index",
            )
        ):
            continue
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _format_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_paragraph_layout(
                        para,
                        alignment=None,
                        first_line_indent_pt=0,
                        line_spacing_pt=TABLE_LINE_SPACING_PT,
                        space_before_pt=0,
                        space_after_pt=0,
                    )
                    for run in para.runs:
                        _set_run_font(run, "宋体", "Times New Roman", TABLE_SIZE_PT, None)


def postprocess_docx(
    path: str | Path,
    metadata: ThesisMetadata,
    report: ConversionReport,
    references: list[ReferenceEntry] | None = None,
) -> None:
    doc = Document(path)
    _format_sections(doc)
    _format_styles(doc)
    _replace_placeholders(doc, report, references=references)
    _replace_citation_placeholders(doc, references, report)
    _format_paragraphs(doc, report)
    equation_numbers = _number_equations(doc, report)
    _replace_equation_reference_placeholders(doc, equation_numbers, report)
    _format_tables(doc)
    _insert_page_section_breaks(doc, report)
    _convert_images_to_top_bottom_wrap(doc, report)
    _keep_figures_with_captions(doc, report)
    if not _replace_marker_with_template(
        doc,
        COVER_MARKER,
        COVER_TEMPLATE_PATH,
        metadata,
        report,
        kind="cover",
        note="官方封面",
    ) and not _replace_marker_with_template(
        doc,
        COVER_MARKER,
        COVER_FALLBACK_TEMPLATE_PATH,
        metadata,
        report,
        kind="cover",
        note="备用官方封面",
    ):
        _replace_cover_placeholder(doc, metadata, report)
    _replace_marker_with_template(
        doc,
        DECLARATION_MARKER,
        DECLARATION_TEMPLATE_PATH,
        metadata,
        report,
        kind="declaration",
        note="原创性声明和授权使用声明",
    )
    _format_sections(doc)
    _apply_section_page_numbering_by_content(doc)
    _format_headers_and_footers(doc, metadata)
    _normalize_pandoc_custom_styles(doc)
    _remove_unrelated_custom_styles(doc)
    _enable_update_fields_on_open(doc)
    doc.save(path)
    report.note("已应用 HEU Engineering A4 页面、正文、标题、题注、表格、参考文献、分页、图片环绕、页眉页脚与目录域后处理。")
