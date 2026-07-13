import base64
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from heuthesis_latex2word.bibliography import ReferenceEntry
from heuthesis_latex2word.bibliography import collect_references
from heuthesis_latex2word.bibliography import parse_bibtex
from heuthesis_latex2word.latex_parser import expand_project, normalize_soft_linebreaks, parse_heusetup
from heuthesis_latex2word.postprocess import (
    _format_paragraphs,
    _format_headers_and_footers,
    _format_sections,
    _format_styles,
    _format_tables,
    _create_reference_numbering,
    _insert_page_section_breaks,
    _keep_figures_with_captions,
    _normalize_pandoc_custom_styles,
    _number_equations,
    _remove_unrelated_custom_styles,
    _replace_equation_reference_placeholders,
    postprocess_docx,
)
from heuthesis_latex2word.preprocess import (
    COVER_MARKER,
    DECLARATION_MARKER,
    GraphicsFallbackReport,
    _prepare_pandoc_math,
    _resolve_graphics_extensions,
    build_declaration_latex,
    make_pandoc_latex,
)
from heuthesis_latex2word.report import ConversionReport


ROOT = Path(__file__).resolve().parents[1]
SAMPLE = ROOT / "HeuThesis_Overleaf" / "examples" / "book" / "bachelor" / "main.tex"


def test_parse_heusetup_nested_values():
    setup = parse_heusetup(
        r"""
        \heusetup{
          ctitlecover={基于~\LaTeX~的哈尔滨工程大学本硕博\\论文模板使用说明},
          cauthor={马冬梅},
          ckeywords={\TeX, \LaTeX, 论文, 模板},
        }
        """
    )
    assert "ctitlecover" in setup
    assert setup["cauthor"] == "马冬梅"
    assert setup["ckeywords"].startswith(r"\TeX")


def test_expand_sample_project_metadata_and_paths():
    project = expand_project(SAMPLE, degree_override="master")
    assert project.metadata.author_cn == "马冬梅"
    assert project.metadata.degree_label == "工程硕士"
    assert project.metadata.secret_level == "公开"
    assert project.metadata.document_number == "no9527"
    assert project.metadata.classified_index == "TM301.2"
    assert project.metadata.udc == "62-5"
    assert project.metadata.author_en == "Ma Dongmei"
    assert project.metadata.affiliation_en == "School of Power and Energy Engineering"
    assert project.metadata.subject_en == "Power Engineering"
    assert "body\\chap01" not in project.latex
    assert "绪论" in project.latex
    assert any(path.name == "figures" for path in project.resource_paths)
    assert len(project.bibliography_files) == 1


def test_preprocess_neutralizes_heu_frontmatter():
    project = expand_project(SAMPLE, degree_override="doctor")
    latex = make_pandoc_latex(project)
    assert r"\usepackage{heuthesis}" not in latex
    assert COVER_MARKER in latex
    assert DECLARATION_MARKER in latex
    assert "HEU_TOC_PLACEHOLDER" in latex
    assert r"\chapter*{结论}" in latex
    assert r"\chapter*{攻读硕士学位期间发表的论文和取得的科研成果}" in latex
    assert r"\chapter*{致谢}" in latex
    assert r"\includegraphics[scale=1.0]{install-texlive.jpg}" in latex
    assert "HEU_CITE_SUPER_" in latex


def test_fallback_declaration_keeps_both_statements_on_one_page():
    latex = build_declaration_latex()
    between_titles = latex.split(r"\section*{学位论文原创性声明}", 1)[1].split(
        r"\section*{学位论文授权使用声明}", 1
    )[0]
    assert r"\newpage" not in between_titles


def test_collect_bibtex_references():
    project = expand_project(SAMPLE, degree_override="master")
    refs = collect_references(project.bibliography_files)
    assert len(refs) >= 10
    assert refs[0].key
    assert refs[0].index == 1
    assert refs[0].text.startswith("[1]")
    assert "武林高手从入门到精通" in refs[0].text
    assert "第 N 次华山论剑" in refs[0].text
    assert " and " not in refs[0].text
    assert "林来兴" in refs[2].text
    assert "空间控制技术" in refs[2].text
    assert all(ref.text.removeprefix(f"[{ref.index}]").strip(" .") for ref in refs)


def test_parse_bibtex_reads_past_first_braced_field(tmp_path):
    bib = tmp_path / "sample.bib"
    bib.write_text(
        """@book{nested,
  language = {zh},
  title = {包含 {嵌套} 花括号的标题},
  author = {甲 and 乙},
  publisher = \"出版社, 总社\",
  year = {2026}
}
""",
        encoding="utf-8",
    )

    entries = parse_bibtex(bib)

    assert len(entries) == 1
    assert entries[0]["language"] == "zh"
    assert entries[0]["title"] == "包含 嵌套 花括号的标题"
    assert entries[0]["author"] == "甲 and 乙"
    assert entries[0]["publisher"] == "出版社, 总社"
    assert entries[0]["year"] == "2026"


def test_eps_only_graphic_gets_generated_png_fallback(tmp_path):
    figures = ROOT / "HeuThesis_Overleaf" / "examples" / "book" / "bachelor" / "figures"
    report = GraphicsFallbackReport()

    resolved = _resolve_graphics_extensions(
        r"\includegraphics[width=4cm]{list}",
        [figures],
        generated_assets_dir=tmp_path / "graphics",
        graphics_report=report,
    )

    assert resolved.endswith(".png}")
    assert len(report.converted_eps) == 1
    source, fallback = report.converted_eps[0]
    assert source.name == "list.eps"
    assert fallback.exists()
    assert fallback.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")
    assert report.unresolved_eps == []


def test_explicit_eps_prefers_existing_compatible_fallback(tmp_path):
    figures = tmp_path / "figures"
    figures.mkdir()
    (figures / "diagram.eps").write_bytes(b"unsupported EPS")
    compatible = figures / "diagram.png"
    compatible.write_bytes(b"\x89PNG\r\n\x1a\nexisting")
    report = GraphicsFallbackReport()

    resolved = _resolve_graphics_extensions(
        r"\includegraphics{diagram.eps}",
        [figures],
        generated_assets_dir=tmp_path / "generated",
        graphics_report=report,
    )

    assert resolved == rf"\includegraphics{{{compatible.resolve().as_posix()}}}"
    assert report.converted_eps == []
    assert report.unresolved_eps == []


def test_unsupported_eps_is_preserved_and_reported(tmp_path):
    figures = tmp_path / "figures"
    figures.mkdir()
    source = figures / "diagram.eps"
    source.write_bytes(b"%!PS-Adobe-3.0 EPSF-3.0\n%%EOF\n")
    report = GraphicsFallbackReport()

    resolved = _resolve_graphics_extensions(
        r"\includegraphics{diagram}",
        [figures],
        generated_assets_dir=tmp_path / "generated",
        graphics_report=report,
    )

    assert resolved == r"\includegraphics{diagram.eps}"
    assert report.converted_eps == []
    assert report.unresolved_eps == [source]
    assert not (tmp_path / "generated").exists()


def test_graphics_resolution_keeps_resource_path_precedence(tmp_path):
    first = tmp_path / "first"
    second = tmp_path / "second"
    first.mkdir()
    second.mkdir()
    (first / "diagram.pdf").write_bytes(b"first resource")
    (second / "diagram.png").write_bytes(b"second resource")

    resolved = _resolve_graphics_extensions(
        r"\includegraphics{diagram}",
        [first, second],
    )

    assert resolved == r"\includegraphics{diagram.pdf}"


def test_reference_numbering_keeps_abstracts_before_instances():
    doc = Document()
    numbering = doc.part.numbering_part.element
    cleanup = OxmlElement("w:numIdMacAtCleanup")
    cleanup.set(qn("w:val"), "1")
    numbering.append(cleanup)
    existing_num_ids = {
        element.get(qn("w:numId")) for element in numbering.findall(qn("w:num"))
    }

    new_num_id = _create_reference_numbering(doc)

    children = list(numbering)
    abstract_positions = [
        index for index, child in enumerate(children) if child.tag == qn("w:abstractNum")
    ]
    instance_positions = [
        index for index, child in enumerate(children) if child.tag == qn("w:num")
    ]
    cleanup_position = children.index(cleanup)
    defined_num_ids = {
        element.get(qn("w:numId")) for element in numbering.findall(qn("w:num"))
    }
    assert max(abstract_positions) < min(instance_positions)
    assert max(instance_positions) < cleanup_position
    assert existing_num_ids < defined_num_ids
    assert str(new_num_id) in defined_num_ids


def test_normalize_soft_linebreaks_keeps_structural_boundaries():
    text = r"""
\section{标题}
第一行正文
第二行正文

\begin{equation}
a=b
c=d
\end{equation}
\item 列表项
下一段
"""
    normalized, count = normalize_soft_linebreaks(text)
    assert "第一行正文 第二行正文" in normalized
    assert "a=b\nc=d" in normalized
    assert "\\item 列表项\n下一段" in normalized
    assert count == 1


def test_prepare_pandoc_math_preserves_code_and_marks_equation_references():
    latex = r"""
普通颜色：\textcolor{red}{红色}。
行内公式 $\textnormal{刚度}=\underleftrightarrow{a}+\underleftarrow{b}+\underrightarrow{c}$，代码 \verb|\underleftrightarrow|。
式(\ref{eq:sample})，另见 \eqref{eq:sample}。
\begin{equation}\label{eq:sample}
  \textcolor[rgb]{0,0,1}{x \Bigm| y}
\end{equation}
\begin{lstlisting}
\textnormal{code}\textcolor[rgb]{0,0,1}{code}\Bigm|
\end{lstlisting}
"""

    prepared = _prepare_pandoc_math(latex)

    assert r"\textcolor{red}{红色}" in prepared
    assert r"\text{刚度}" in prepared
    assert r"\underset{\leftrightarrow}{a}" in prepared
    assert r"\underset{\leftarrow}{b}" in prepared
    assert r"\underset{\rightarrow}{c}" in prepared
    assert r"\verb|\underleftrightarrow|" in prepared
    assert r"x \mid" in prepared
    assert r"\textcolor[rgb]{0,0,1}{x" not in prepared
    assert "HEU_EQUATION_MARK_" in prepared
    assert "HEU_EQREF_PLAIN_" in prepared
    assert "HEU_EQREF_PAREN_" in prepared
    marker_end = prepared.index("_END", prepared.index("HEU_EQUATION_MARK_"))
    assert prepared[marker_end + 4 :].lstrip("\n").startswith(r"\begin{lstlisting}")
    assert prepared[marker_end + 4 :].startswith("\n\n")
    assert r"\textnormal{code}\textcolor[rgb]{0,0,1}{code}\Bigm|" in prepared


def test_postprocess_adds_cover_styles_bookmarks_and_ref_fields(tmp_path):
    project = expand_project(SAMPLE, degree_override="master")
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("HEU_COVER_PLACEHOLDER")
    doc.add_paragraph(DECLARATION_MARKER)
    toc_title = doc.add_paragraph("目录")
    toc_title.style = "Heading 1"
    doc.add_paragraph("HEU_TOC_PLACEHOLDER")
    doc.add_paragraph(
        "正文 HEU_CITE_SUPER_RFhNMjAwNQ.ZXhhbXBsZQ_END；"
        "行文 HEU_CITE_TEXT_ZXhhbXBsZQ.RFhNMjAwNQ_END；"
        "连续 HEU_CITE_SUPER_RFhNMjAwNQ_ENDHEU_CITE_SUPER_ZXhhbXBsZQ_END"
    )
    doc.add_paragraph("HEU_REFERENCES_PLACEHOLDER")
    doc.save(doc_path)

    postprocess_docx(
        doc_path,
        project.metadata,
        ConversionReport(),
        references=[
            ReferenceEntry(key="DXM2005", index=1, text="[1] First reference."),
            ReferenceEntry(key="example", index=2, text="[2] Second reference."),
        ],
    )

    processed = Document(doc_path)
    style_names = {style.name for style in processed.styles}
    xml = processed._element.xml
    assert "封面题名" in style_names
    assert "参考文献" in style_names
    assert "封面宋体小2" in style_names
    assert "封面2号黑体英文论文名称" in style_names
    assert 'w:name="HEU_REF_DXM2005"' in xml
    assert r' REF HEU_REF_DXM2005 \n \# "0" \h ' in xml
    assert r' REF HEU_REF_example \n \# "0" \h ' in xml
    assert 'TOC \\o "1-3" \\h \\z \\u' in xml
    assert xml.count('w:dirty="true"') >= 7
    assert 'w:vertAlign w:val="superscript"' in xml

    citation = next(para for para in processed.paragraphs if para.text.startswith("正文 ["))
    assert citation.text == "正文 [1,2]；行文 [1,2]；连续 [1][2]"
    reference_paragraphs = [
        para for para in processed.paragraphs if para.style.name == "参考文献"
    ]
    assert [para.text for para in reference_paragraphs] == [
        "First reference.",
        "Second reference.",
    ]
    num_ids = {
        para._p.pPr.numPr.numId.get(qn("w:val"))
        for para in reference_paragraphs
    }
    assert len(num_ids) == 1
    assert 'w:lvlText w:val="[%1]"' in processed.part.numbering_part.element.xml
    toc_title = next(para for para in processed.paragraphs if para.text == "目录")
    outline = toc_title._p.pPr.find(qn("w:outlineLvl"))
    assert outline is not None
    assert outline.get(qn("w:val")) == "9"
    update_fields = processed.settings._element.find(qn("w:updateFields"))
    assert update_fields is not None
    assert update_fields.get(qn("w:val")) == "true"

    all_paragraphs = processed.paragraphs
    declaration = next(para for para in all_paragraphs if para.text == "学位论文原创性声明")
    authorization = next(para for para in all_paragraphs if para.text == "学位论文授权使用声明")
    body = list(processed._element.body)
    between_statements = body[body.index(declaration._p) + 1 : body.index(authorization._p)]
    page_breaks = [
        br
        for element in between_statements
        for br in element.findall(".//" + qn("w:br"))
        if br.get(qn("w:type")) == "page"
    ]
    assert not page_breaks
    authorization_index = all_paragraphs.index(authorization)
    authorization_signature = next(
        para
        for para in all_paragraphs[authorization_index + 1 :]
        if para.text.strip().startswith("作者(签字)") and "导师(签字)" in para.text
    )
    authorization_date = next(
        para
        for para in all_paragraphs[authorization_index + 1 :]
        if para.text.strip().startswith("日期") and para.text.count("日期") == 2
    )
    assert authorization_signature.text == "作者(签字)：\t导师(签字)："
    assert authorization_signature.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert authorization_date.text.count("\t") == 1
    assert len(authorization_date._p.findall(".//" + qn("w:r") + "/" + qn("w:tab"))) == 1
    assert authorization_date._p.pPr.find(qn("w:tabs")) is not None
    assert not authorization_date._p.findall(".//" + qn("w:br"))
    for line in (authorization_signature, authorization_date):
        ppr = line._p.pPr
        assert ppr.find(qn("w:numPr")) is None
        for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore"):
            assert ppr.find(qn(tag)) is None
        indent = ppr.find(qn("w:ind"))
        for attr in ("w:left", "w:right", "w:firstLine"):
            assert int(indent.get(qn(attr), "0")) == 0


def test_body_odd_header_uses_chapter_and_even_header_uses_school_line():
    project = expand_project(SAMPLE, degree_override="master")
    doc = Document()
    doc.add_paragraph("封面")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("摘要")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("第1章 绪论")

    abstract_section = doc.sections[1]
    body_section = doc.sections[2]
    roman_numbering = OxmlElement("w:pgNumType")
    roman_numbering.set(qn("w:fmt"), "upperRoman")
    roman_numbering.set(qn("w:start"), "1")
    abstract_section._sectPr.append(roman_numbering)
    page_numbering = OxmlElement("w:pgNumType")
    page_numbering.set(qn("w:fmt"), "decimal")
    page_numbering.set(qn("w:start"), "1")
    body_section._sectPr.append(page_numbering)

    _format_sections(doc)
    _format_headers_and_footers(doc, project.metadata)
    _format_headers_and_footers(doc, project.metadata)

    school_line = "哈尔滨工程大学工程硕士学位论文"
    assert doc.settings.odd_and_even_pages_header_footer is True
    assert doc.settings._element.find(qn("w:evenAndOddHeaders")) is not None
    abstract_odd = abstract_section.header.paragraphs[0]
    assert abstract_odd.text == school_line
    assert "STYLEREF" not in abstract_section.header._element.xml
    assert abstract_section.even_page_header.paragraphs[0].text == school_line
    assert abstract_odd.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert abstract_odd.paragraph_format.first_line_indent.pt == 0
    assert abstract_odd.paragraph_format.left_indent.pt == 0
    assert abstract_odd.paragraph_format.right_indent.pt == 0
    for header_paragraph in (
        abstract_odd,
        abstract_section.even_page_header.paragraphs[0],
        body_section.header.paragraphs[0],
        body_section.even_page_header.paragraphs[0],
    ):
        borders = header_paragraph._p.pPr.find(qn("w:pBdr"))
        assert borders is not None
        assert len(header_paragraph._p.pPr.findall(qn("w:pBdr"))) == 1
        bottom = borders.find(qn("w:bottom"))
        assert bottom.get(qn("w:val")) == "thickThinSmallGap"
        assert bottom.get(qn("w:sz")) == "24"
        assert bottom.get(qn("w:space")) == "1"
        assert bottom.get(qn("w:color")) == "auto"
        for edge in ("top", "left", "right"):
            side = borders.find(qn(f"w:{edge}"))
            assert side.get(qn("w:val")) == "none"
            assert side.get(qn("w:sz")) == "0"
            assert side.get(qn("w:space")) == "0"
            assert side.get(qn("w:color")) == "auto"

    for header in (
        doc.sections[0].header,
        doc.sections[0].even_page_header,
        doc.sections[0].first_page_header,
    ):
        assert header.paragraphs[0]._p.pPr.find(qn("w:pBdr")) is None
    for section in doc.sections:
        for footer in (section.footer, section.even_page_footer, section.first_page_footer):
            assert footer.paragraphs[0]._p.pPr.find(qn("w:pBdr")) is None

    odd_xml = body_section.header._element.xml
    assert 'STYLEREF "章标题" \\* MERGEFORMAT' in odd_xml
    assert body_section.header.paragraphs[0].text == "第1章 绪论"
    assert body_section.header.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_section.header.paragraphs[0].paragraph_format.first_line_indent.pt == 0
    assert body_section.even_page_header.paragraphs[0].text == school_line
    assert "PAGE" in body_section.footer._element.xml
    assert "PAGE" in body_section.even_page_footer._element.xml


def test_postprocess_uses_styles_for_text_and_keeps_image_caption_tight():
    doc = Document()
    heading = doc.add_paragraph("样式标题")
    heading.style = "Heading 1"
    section_one = doc.add_paragraph("第一节")
    section_one.style = "Heading 2"
    subsection_one = doc.add_paragraph("第一小节")
    subsection_one.style = "Heading 3"
    subsection_two = doc.add_paragraph("第二小节")
    subsection_two.style = "Heading 3"
    section_two = doc.add_paragraph("第二节")
    section_two.style = "Heading 2"
    subsection_three = doc.add_paragraph("第一小节")
    subsection_three.style = "Heading 3"
    summary = doc.add_paragraph("本章小结")
    summary.style = "Heading 2"
    body = doc.add_paragraph("样式正文")
    image = doc.add_paragraph()
    image.add_run().add_picture(str(ROOT / "HeuThesis_Overleaf" / "heulogo.jpg"))
    caption = doc.add_paragraph("校徽")
    caption.style = "Caption"

    _format_styles(doc)
    _format_paragraphs(doc)

    assert heading.style.name == "章标题"
    assert heading.text == "第1章 样式标题"
    assert section_one.text == "1.1 第一节"
    assert subsection_one.text == "1.1.1 第一小节"
    assert subsection_two.text == "1.1.2 第二小节"
    assert section_two.text == "1.2 第二节"
    assert subsection_three.text == "1.2.1 第一小节"
    assert summary.text == "1.3 本章小结"
    assert section_one.style.name == "节标题"
    assert subsection_one.style.name == "小节标题"
    assert body.style.name == "正文"
    assert caption.style.name == "图表题注"
    assert image.style.name == "图片"
    assert image.style.paragraph_format.space_after.pt == 0
    assert image.style.paragraph_format.line_spacing.pt == 1.0
    assert image.style.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert doc.styles["正文"]._element.find(qn("w:qFormat")) is not None
    assert doc.styles["章标题"]._element.find(qn("w:qFormat")) is not None
    assert doc.styles["图表题注"]._element.find(qn("w:qFormat")) is not None
    assert doc.styles["toc 1"].style_id == "TOC1"
    assert doc.styles["toc 1"].font.name == "Times New Roman"
    assert doc.styles["toc 1"].paragraph_format.left_indent.pt == 0
    assert doc.styles["toc 2"].paragraph_format.left_indent.pt == 24
    assert doc.styles["toc 3"].paragraph_format.left_indent.pt == 48
    assert qn("w:customStyle") not in doc.styles["toc 1"]._element.attrib
    assert heading.runs[0].font.size is None
    assert body.runs[0].font.size is None


def test_numbered_equation_has_chapter_sequence_tabs_bookmark_and_ref():
    label = "eq:sample"
    token = base64.urlsafe_b64encode(label.encode("utf-8")).decode("ascii").rstrip("=")
    doc = Document()
    heading = doc.add_paragraph("公式")
    heading.style = "Heading 1"
    equation = doc.add_paragraph()
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=y"
    math_run.append(math_text)
    math.append(math_run)
    math_para.append(math)
    equation._p.append(math_para)
    marker = doc.add_paragraph(f"HEU_EQUATION_MARK_{token}_END\u00a0◻")
    reference = doc.add_paragraph(f"如式(HEU_EQREF_PLAIN_{token}_END)所示。")

    _format_styles(doc)
    _format_paragraphs(doc)
    equation_numbers = _number_equations(doc)
    _replace_equation_reference_placeholders(doc, equation_numbers)

    xml = doc._element.xml
    assert marker._p.getparent() is None
    assert equation_numbers == {label: "1-1"}
    assert "HEU_EQUATION_MARK_" not in xml
    assert "HEU_EQREF_" not in xml
    assert "SEQ HEUEquation1" in xml
    assert "REF HEU_EQ_eq_sample" in xml
    assert 'w:name="HEU_EQ_eq_sample"' in xml
    assert equation.text == "\t\t(1-1)"
    assert equation._p.find(qn("m:oMath")) is not None
    assert reference.text == "如式(1-1)所示。"
    tabs = equation._p.pPr.find(qn("w:tabs"))
    assert [(tab.get(qn("w:val")), tab.get(qn("w:pos"))) for tab in tabs] == [
        ("center", "4536"),
        ("right", "9072"),
    ]


def test_figure_caption_and_publication_heading_keep_with_following_content():
    doc = Document()
    achievement = doc.add_paragraph("攻读硕士学位期间发表的论文和取得的科研成果")
    achievement.style = "Heading 1"
    publication = doc.add_paragraph("（一）发表的相关论文")
    publication.style = "Heading 2"
    doc.add_paragraph("论文条目")
    image = doc.add_paragraph()
    image.add_run().add_picture(str(ROOT / "HeuThesis_Overleaf" / "heulogo.jpg"))
    caption = doc.add_paragraph("校徽")
    caption.style = "Caption"

    _format_styles(doc)
    _format_paragraphs(doc)
    _insert_page_section_breaks(doc)
    _keep_figures_with_captions(doc)

    assert image._p.pPr.find(qn("w:keepNext")) is not None
    assert image._p.pPr.find(qn("w:keepLines")) is not None
    assert caption._p.pPr.find(qn("w:keepLines")) is not None
    assert publication.style.paragraph_format.keep_with_next is True
    previous = publication._p.getprevious()
    assert previous is achievement._p
    assert previous.find(qn("w:pPr") + "/" + qn("w:sectPr")) is None


def test_front_matter_and_chapters_get_required_section_page_numbering():
    doc = Document()
    doc.add_paragraph("封面占位")
    abstract = doc.add_paragraph("摘要")
    abstract.style = "Heading 1"
    doc.add_paragraph("摘要正文")
    toc = doc.add_paragraph("目录")
    toc.style = "Heading 1"
    chapter = doc.add_paragraph("绪论")
    chapter.style = "Heading 1"

    _format_styles(doc)
    _format_paragraphs(doc)
    _insert_page_section_breaks(doc)

    assert abstract.text == "摘要"
    assert toc.text == "目录"
    assert chapter.text == "第1章 绪论"

    for front_title in (abstract, toc):
        outline = front_title._p.pPr.find(qn("w:outlineLvl"))
        assert outline is not None
        assert outline.get(qn("w:val")) == "9"

    xml = doc._element.xml
    assert 'w:type w:val="nextPage"' in xml
    assert 'w:type w:val="oddPage"' in xml
    assert 'w:pgNumType w:fmt="upperRoman" w:start="1"' in xml
    assert 'w:pgNumType w:fmt="decimal" w:start="1"' in xml


def test_inline_math_body_uses_exact_body_line_spacing_style():
    doc = Document()
    para = doc.add_paragraph("含公式的正文")
    para._p.append(OxmlElement("m:oMath"))

    _format_styles(doc)
    _format_paragraphs(doc)

    assert para.style.name == "正文（内嵌公式）"
    assert para.style.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert para.style.paragraph_format.line_spacing.pt == 20.5


def test_image_table_borders_are_removed():
    doc = Document()
    _format_styles(doc)
    table = doc.add_table(rows=1, cols=2)
    for cell in table.rows[0].cells:
        cell.paragraphs[0].add_run().add_picture(str(ROOT / "HeuThesis_Overleaf" / "heulogo.jpg"))

    _format_tables(doc)

    xml = table._tbl.xml
    assert table.style.name == "FigureTable"
    assert 'w:val="single"' not in xml
    assert xml.count('w:val="nil"') >= 6
    assert {cell.paragraphs[0].style.name for cell in table.rows[0].cells} == {"图片"}


def test_empty_figure_layout_tables_are_removed_without_touching_regular_empty_tables():
    doc = Document()
    _format_styles(doc)
    regular_empty = doc.add_table(rows=1, cols=1)
    doc.add_paragraph("普通表格后的正文")
    caption = doc.add_paragraph("图 1.1 示例")
    caption.style = "图表题注"
    empty_figure = doc.add_table(rows=1, cols=2)
    populated_figure = doc.add_table(rows=1, cols=2)
    populated_figure.cell(0, 0).text = "(a) 子图"
    trailing_caption = doc.add_paragraph("图 1.2 组合图")
    trailing_caption.style = "图表题注"

    _format_tables(doc)

    assert regular_empty._tbl.getparent() is not None
    assert empty_figure._tbl.getparent() is None
    assert populated_figure._tbl.getparent() is not None
    assert populated_figure.style.name == "FigureTable"
    assert 'w:val="single"' not in populated_figure._tbl.xml


def test_unrelated_custom_styles_are_removed_and_legacy_heu_styles_hidden():
    doc = Document()
    _format_styles(doc)
    used = doc.styles.add_style("Temporary Used Style", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Temporary Unused Style", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("HEU Body", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("kept", style=used)

    _remove_unrelated_custom_styles(doc)

    style_names = {style.name for style in doc.styles}
    assert "Temporary Used Style" in style_names
    assert "Temporary Unused Style" not in style_names
    assert "HEU Body" not in style_names


def test_pandoc_custom_styles_are_normalized_before_cleanup():
    doc = Document()
    _format_styles(doc)
    doc.styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Verbatim Char", WD_STYLE_TYPE.CHARACTER)
    para = doc.add_paragraph("compact", style="Compact")
    code = doc.add_paragraph("command: ")
    code.add_run("pytest").style = "Verbatim Char"

    _normalize_pandoc_custom_styles(doc)
    _remove_unrelated_custom_styles(doc)

    style_names = {style.name for style in doc.styles}
    assert para.style.name == "正文"
    assert code.runs[0].style.name == "Default Paragraph Font"
    assert "Compact" not in style_names
    assert "Verbatim Char" not in style_names
