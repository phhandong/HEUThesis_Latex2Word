from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from heuthesis_latex2word.latex_parser import ThesisMetadata
from heuthesis_latex2word.postprocess import (
    COVER_TEMPLATE_PATH,
    _apply_cover_metadata,
    _copy_missing_template_styles,
)


def test_cover_template_path_uses_repo_root_template():
    assert COVER_TEMPLATE_PATH.name == "cover.docx"
    assert COVER_TEMPLATE_PATH.parent.name == "HEUThesis_Latex2Word"


def _metadata() -> ThesisMetadata:
    return ThesisMetadata(
        degree="master",
        classified_index="TM301.2",
        state_secrets="公开",
        udc="62-5",
        document_number="no9527",
        title_cover_cn="基于 FPGA 的高能效 XXX 应用系统",
        title_en="An Energy Efficient FPGA-based System\nfor XXX Applications",
        author_cn="马冬梅",
        supervisor_cn="孔夫子 教授",
        cosupervisor_cn="李四 高级工程师",
        subject_cn="电子信息",
        affiliation_cn="水声工程学院",
        submit_date_cn="2026年6月",
        oral_date_cn="2026年7月",
        author_en="Ma Dongmei",
        supervisor_en="Prof. Kong Fuzi",
        cosupervisor_en="Li Si",
        subject_en="Electronic Information",
        affiliation_en="College of Underwater Acoustic Engineering",
        submit_date_en="June, 2026",
        oral_date_en="July, 2026",
    )


def test_cover_metadata_removes_annotations_and_aligns_cover_fields():
    doc = Document(COVER_TEMPLATE_PATH)
    _apply_cover_metadata(doc, _metadata())

    visible_text = "\n".join(p.text for p in doc.paragraphs)
    visible_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    for annotation in ("宋体小四号", "宋体小二号", "黑体二号", "楷体小二号", "Times New Roman四号"):
        assert annotation not in visible_text

    top_line = doc.paragraphs[0]
    second_line = doc.paragraphs[1]
    for paragraph in (top_line, second_line):
        tab_stops = list(paragraph.paragraph_format.tab_stops)
        assert "\t" in paragraph.text
        assert tab_stops
        assert tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT

    assert "分类号：TM301.2" in top_line.text
    assert "密级：公开" in top_line.text
    assert any(run.underline and "TM301.2" in run.text for run in top_line.runs)
    assert any(run.underline and "公开" in run.text for run in top_line.runs)

    info_lines = doc.paragraphs[12:16]
    assert [p.text for p in info_lines] == [
        "　　　　　　　　　硕士研究生：　马冬梅",
        "　　　　　　　　　指导教师：　　孔夫子 教授",
        "　　　　　　　　　校外导师：　　李四 高级工程师",
        "　　　　　　　　　学位类别：　　工程硕士",
    ]
    assert {p.alignment for p in info_lines} == {WD_ALIGN_PARAGRAPH.JUSTIFY}
    assert {p.paragraph_format.left_indent.cm for p in info_lines} == {0}
    assert all("firstLineChars" not in p._p.xml for p in info_lines)
    assert all(not list(p.paragraph_format.tab_stops) for p in info_lines)

    assert doc.tables[0].rows[7].cells[1].text == "：哈尔滨工程大学"
    assert doc.tables[1].rows[0].cells[1].text == "Ma Dongmei"


def test_english_cover_uses_cover_template_styles_and_fields():
    doc = Document(COVER_TEMPLATE_PATH)
    for para in doc.paragraphs:
        if para.style and para.style.name == "封面2号黑体英文论文名称":
            para.runs[0].text = "English title placeholder changed"
            for run in para.runs[1:]:
                run.text = ""

    _apply_cover_metadata(doc, _metadata())

    assert doc.paragraphs[39].text == "Classified Index: TM301.2"
    assert doc.paragraphs[40].text == "U.D.C: 62-5"
    assert doc.paragraphs[44].text == "A Thesis for the Degree of Master of Engineering"
    assert doc.paragraphs[47].text == "An Energy Efficient FPGA-based System"
    assert doc.paragraphs[48].text == "for XXX Applications"
    assert doc.tables[1].rows[0].cells[1].text == "Ma Dongmei"
    assert doc.tables[1].rows[1].cells[1].text == "Prof. Kong Fuzi"
    assert doc.tables[1].rows[2].cells[1].text == "Li Si"
    assert doc.tables[1].rows[3].cells[1].text == "Electronic Information"
    assert doc.tables[1].rows[4].cells[1].text == "College of Underwater Acoustic Engineering"
    assert doc.tables[1].rows[5].cells[1].text == "June, 2026"
    assert doc.tables[1].rows[6].cells[1].text == "July, 2026"
    assert doc.tables[1].rows[7].cells[1].text == "Harbin Engineering University"


def test_cover_template_custom_styles_are_copied_to_output_doc():
    target = Document()
    template = Document(COVER_TEMPLATE_PATH)

    _copy_missing_template_styles(target, template)

    style_names = {style.name for style in target.styles}
    assert "封面宋体小2" in style_names
    assert "封面2号黑体英文论文名称" in style_names
    assert target.styles["封面宋体小2"].font.name == "Times New Roman"
    assert target.styles["封面2号黑体英文论文名称"].font.name == "Times New Roman"
